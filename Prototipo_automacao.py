from docx import Document
from docx.shared import Inches
import io
import os
import shutil
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from tkinter.scrolledtext import ScrolledText
import pandas as pd
from docx import Document as DocxDocument
from datetime import datetime, timedelta
import threading
import logging
import json
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import zipfile
from pdf2docx import Converter
from docx2pdf import convert
import pythoncom
import comtypes.client
from difflib import get_close_matches
import locale
from PIL import Image, ImageTk
import tempfile
import queue
import sys
import traceback
import gc
from concurrent.futures import ThreadPoolExecutor, as_completed
import psutil
import subprocess
import platform
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import webbrowser
import time
import winsound
import schedule
import pytesseract
import csv
import sqlite3
from collections import deque
import re
import shlex
import unicodedata
import copy
import glob
from typing import List, Dict 
from datetime import datetime
from typing import Dict, List
from concurrent.futures import ThreadPoolExecutor




# Configuração de localização (locale)
locale.setlocale(locale.LC_ALL, '')  
# Define a localidade padrão do sistema operacional para todas as categorias.
# Isso influencia, por exemplo, em como datas, moedas e números são formatados.

_ = lambda x: x  # Função de tradução fictícia (placeholder)
# Essa função é usada como um substituto para tradução (normalmente gettext).
# Aqui, ela simplesmente retorna o mesmo texto sem traduzir.
# Em um futuro suporte a múltiplos idiomas, isso pode ser substituído por uma função real de tradução.

# Configuração aprimorada de logs
logging.basicConfig(
    level=logging.DEBUG,  # Nível de log alterado de INFO para DEBUG (mais detalhado)
    format='%(asctime)s - %(levelname)s - %(message)s',  # Formato da mensagem de log com data/hora, nível e mensagem
    handlers=[
        logging.FileHandler('document_converter.log', encoding='utf-8'),  # Salva os logs em um arquivo
        logging.StreamHandler()  # Também exibe os logs no console
    ]
)
logger = logging.getLogger(__name__)  # Obtém um logger nomeado com o nome do módulo atual


class LibreOfficeDocument:
    """
    Classe para manipular documentos via LibreOffice em modo headless.

    Permite abrir, converter e salvar documentos com controle de timeout e
    tratamento de erros aprimorado, suportando operações em segundo plano.
    
    Atributos:
        libreoffice_path (str): Caminho para o executável do LibreOffice.
        temp_dir (str): Diretório temporário para arquivos intermediários.
        timeout (int): Tempo limite padrão para operações com subprocessos.
    """
    
    def __init__(self, libreoffice_path: str, temp_dir: str = None, timeout: int = 60):
        if not os.path.exists(libreoffice_path):
            raise FileNotFoundError(f"LibreOffice executable not found at: {libreoffice_path}")
        self.libreoffice_path = libreoffice_path
        self.temp_dir = temp_dir or tempfile.mkdtemp()
        self.document_path = None
        self.process = None
        self.timeout = timeout  # Timeout padrão de 60 segundos
        self._soffice_processes = []  # Para rastrear todos os processos do LibreOffice

    def _run_with_timeout(self, cmd: List[str], timeout: int = None) -> subprocess.CompletedProcess:
        """Executa um comando com timeout"""
        timeout = timeout or self.timeout
        try:
            return subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=timeout,
                check=True,
                encoding='utf-8',
                errors='ignore'
            )
        except subprocess.TimeoutExpired:
            self._kill_libreoffice_processes()
            raise RuntimeError(f"LibreOffice operation timed out after {timeout} seconds")
        except subprocess.CalledProcessError as e:
            error_msg = self._parse_libreoffice_error(e.stderr)
            raise RuntimeError(f"LibreOffice conversion failed: {error_msg}") from e

    def _kill_libreoffice_processes(self):
        """Mata todos os processos do LibreOffice relacionados"""
        for proc in self._soffice_processes:
            try:
                proc.terminate()
                proc.wait(timeout=5)
            except:
                pass
        self._soffice_processes = []

    def _parse_libreoffice_error(self, error_output: str) -> str:
        """Analisa a saída de erro do LibreOffice para mensagens mais amigáveis"""
        common_errors = {
            "could not find an office installation": "LibreOffice não está instalado corretamente",
            "file could not be loaded": "O arquivo não pôde ser carregado - pode estar corrompido",
            "password protected": "O arquivo está protegido por senha",
            "out of memory": "Memória insuficiente para a operação",
            "filter not found": "Filtro de conversão não disponível"
        }
        
        # Busca por erros conhecidos
        for error, message in common_errors.items():
            if error.lower() in error_output.lower():
                return message
        
        # Retorna as primeiras linhas do erro se não reconhecer
        return "\n".join(error_output.splitlines()[:3])
    
    def check_libreoffice_available(self) -> bool:
        """Verifica se o LibreOffice está disponível e funcionando"""
        try:
            # Teste simples de versão
            result = self._run_with_timeout([self.libreoffice_path, '--version'], timeout=10)
            
            # Verifica se há saída válida
            if not result.stdout.strip():
                return False
                
            # Verifica se há processos zumbis do LibreOffice
            self._clean_zombie_processes()
            
            return True
            
        except Exception as e:
            logger.error(f"LibreOffice check failed: {str(e)}")
            return False

    def _clean_zombie_processes(self):
        """Limpa processos do LibreOffice que possam estar travados"""
        if platform.system() == "Windows":
            self._clean_windows_processes()
        else:
            self._clean_unix_processes()

    def _clean_windows_processes(self):
        """Limpeza específica para Windows"""
        try:
            # Usa tasklist para encontrar processos soffice
            result = subprocess.run(
                ['tasklist', '/FI', 'IMAGENAME eq soffice*'],
                capture_output=True, text=True
            )
            
            if "soffice" in result.stdout:
                # Mata todos os processos soffice
                subprocess.run(['taskkill', '/F', '/IM', 'soffice*'], check=True)
                
        except Exception as e:
            logger.warning(f"Could not clean Windows LibreOffice processes: {str(e)}")

    def _clean_unix_processes(self):
        """Limpeza específica para Unix/Linux"""
        try:
            # Usa ps e grep para encontrar processos
            subprocess.run(
                ['pkill', '-f', 'soffice'],
                stderr=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL
            )
        except Exception as e:
            logger.warning(f"Could not clean Unix LibreOffice processes: {str(e)}")
    
    def convert_to(self, output_path: str, target_format: str, timeout: int = None) -> str:
        """Converte para um formato especificado com timeout"""
        if not self.document_path:
            raise ValueError("No document is open")
        
        output_folder = os.path.dirname(output_path)
        os.makedirs(output_folder, exist_ok=True)
        
        cmd = [
            self.libreoffice_path,
            '--headless',
            '--convert-to', target_format,
            '--outdir', output_folder,
            self.document_path
        ]
        
        try:
            result = self._run_with_timeout(cmd, timeout)
            
            # Verifica se o arquivo de saída foi criado
            expected_ext = f".{target_format.lower().split(':')[0]}"
            output_file = os.path.join(
                output_folder, 
                os.path.splitext(os.path.basename(self.document_path))[0] + expected_ext
            )
            
            if not os.path.exists(output_file):
                self._kill_libreoffice_processes()
                raise RuntimeError(f"Conversion succeeded but output file not found at: {output_file}")

            
            return output_file
            
        except Exception as e:
            self._kill_libreoffice_processes()
            logger.error(f"Conversion error: {str(e)}")
            raise


    def open(self, file_path: str, timeout: int = None) -> 'LibreOfficeDocument':
        """Abre um documento no LibreOffice em modo headless com timeout"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        cmd = [
            self.libreoffice_path,
            '--headless',
            '--norestore',
            '--nologo',
            '--nodefault',
            '--nofirststartwizard',
            file_path
        ]
        
        try:
            # Inicia o processo e armazena para possível finalização
            self.process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                encoding='utf-8',
                errors='ignore'
            )
            self._soffice_processes.append(self.process)
            
            # Verifica se o processo está rodando após um breve delay
            time.sleep(0.5)
            if self.process.poll() is not None:
                err = self.process.stderr.read() if self.process.stderr else "Unknown error"
                raise RuntimeError(f"LibreOffice failed to start: {err}")
                
            self.document_path = file_path
            return self
            
        except Exception as e:
            self._kill_libreoffice_processes()
            logger.error(f"Error opening document: {str(e)}")
            raise
    
    def save_as(self, output_path: str):
        """Salva o documento em um novo formato"""
        if not self.document_path:
            raise ValueError("No document is open")
        
        # Normaliza os caminhos para evitar problemas com barras
        docx_path = os.path.normpath(self.document_path)
        output_folder = os.path.normpath(os.path.dirname(output_path))
        
        cmd = [
            self.libreoffice_path,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_folder,
            docx_path
        ]
        
        self.log_message(f"[DEBUG] Executável LibreOffice: {self.libreoffice_path}")
        self.log_message(f"[DEBUG] Comando completo: {' '.join(cmd)}")

        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=30
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")
        
        # Retorna o caminho completo do arquivo PDF gerado
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        return os.path.join(output_folder, pdf_filename)
        
    def close(self):
        """Fecha o documento e limpa recursos de forma robusta"""
        attempts = 3
        while attempts > 0:
            try:
                if self.process:
                    self.process.terminate()
                    try:
                        self.process.wait(timeout=5)
                    except subprocess.TimeoutExpired:
                        self.process.kill()
                
                self._kill_libreoffice_processes()
                
                if os.path.exists(self.temp_dir):
                    shutil.rmtree(self.temp_dir, ignore_errors=True)
                
                break
            except Exception as e:
                attempts -= 1
                if attempts == 0:
                    logger.error(f"Failed to clean up resources: {str(e)}")
                time.sleep(1)

    def __enter__(self):
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

class DocumentManager:
    """Gerencia documentos usando o backend disponível (Word ou LibreOffice)"""
    
    def __init__(self, use_libreoffice: bool = False, libreoffice_path: str = None):
        self.use_libreoffice = use_libreoffice
        self.libreoffice_path = libreoffice_path
    
    def open_document(self, file_path: str):
        """Abre um documento usando o backend apropriado"""
        if self.use_libreoffice and self.libreoffice_path:
            return LibreOfficeDocument(self.libreoffice_path).open(file_path)
        else:
            try:
                return DocxDocument(file_path)

            except Exception as e:
                raise RuntimeError(f"Failed to open document with Word: {str(e)}")
    
    def save_document(self, doc, output_path: str):
        """Salva um documento usando o backend apropriado"""
        try:
            # Verificação mais robusta do tipo de documento
            if hasattr(doc, 'save') and callable(doc.save):  # Para DocxDocument
                doc.save(output_path)
                return output_path
            elif hasattr(doc, 'save_as') and callable(doc.save_as):  # Para LibreOfficeDocument
                return doc.save_as(output_path)
            else:
                raise ValueError("Unsupported document type - missing save method")
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise

class ExecutionHistoryManager:
    """Manages execution history in a SQLite database"""
    def __init__(self):
        self.db_path = "execution_history.db"
        self._init_db()
        
    def _init_db(self):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS execution_history (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    operation_type TEXT,
                    input_files TEXT,
                    output_folder TEXT,
                    status TEXT,
                    processed_count INTEGER,
                    error_count INTEGER,
                    duration REAL,
                    log_text TEXT
                )
            """)
            conn.commit()
    
    def add_record(self, operation_type: str, input_files: str, output_folder: str,
                  status: str, processed_count: int, error_count: int,
                  duration: float, log_text: str):
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO execution_history (
                    operation_type, input_files, output_folder,
                    status, processed_count, error_count,
                    duration, log_text
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                operation_type, input_files, output_folder,
                status, processed_count, error_count,
                duration, log_text
            ))
            conn.commit()
    
    def get_recent_history(self, limit=20) -> List[Dict]:
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("""
                SELECT * FROM execution_history 
                ORDER BY timestamp DESC 
                LIMIT ?
            """, (limit,))
            return [dict(row) for row in cursor.fetchall()]

class EnhancedConfigManager:
    """Manages persistent application configurations with versioning"""
    CONFIG_FILE = "config_v4.json"
    CONFIG_VERSION = 4
    
    def __init__(self):
        self.config = self._load_config()
        self.history = ExecutionHistoryManager()
        
    def _load_config(self) -> Dict:
        """Loads configuration from file or returns default with migration if needed"""
        default_config = {
            "version": self.CONFIG_VERSION,
            "recent_files": {
                "excel": "",
                "template": "",
                "output": "",
                "docx": [],
                "docx_output": "",
                "pdf": [],
                "pdf_output": ""
            },
            "window_size": "900x750",
            "zip_output": False,
            "theme": "clam",
            "max_threads": min(4, os.cpu_count() or 1),
            "last_tab": 0,
            "column_mappings": {},
            "use_libreoffice": False,
            "libreoffice_path": self._detect_libreoffice_path(),
            "progress_throttle": 0.1,
            "enable_sounds": True,
            "show_notifications": True,
            "auto_load_last_files": True,
            "ocr_enabled": False,
            "tesseract_path": self._detect_tesseract_path(),
            "max_history_items": 50
        }
        
        try:
            if os.path.exists(self.CONFIG_FILE):
                with open(self.CONFIG_FILE, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)
                    
                    # Migration from old versions
                    if loaded_config.get("version", 1) < self.CONFIG_VERSION:
                        loaded_config = self._migrate_config(loaded_config)
                    
                    # Merge with default config for new keys
                    return {**default_config, **loaded_config}
        except Exception as e:
            logger.error(f"Error loading config: {str(e)}")
            logger.debug(traceback.format_exc())
            
        return default_config
    
    def _migrate_config(self, old_config: Dict) -> Dict:
        """Migrates configurations from old versions"""
        if old_config.get("version", 1) == 1:
            new_config = old_config.copy()
            new_config["version"] = 2
            new_config["recent_files"]["docx"] = [new_config["recent_files"].get("docx", "")]
            new_config["recent_files"]["pdf"] = [new_config["recent_files"].get("pdf", "")]
            new_config["max_threads"] = 4
            new_config["last_tab"] = 0
            return new_config
        elif old_config.get("version", 2) == 2:
            new_config = old_config.copy()
            new_config["version"] = 3
            new_config["column_mappings"] = {}
            new_config["use_libreoffice"] = False
            new_config["libreoffice_path"] = self._detect_libreoffice_path()
            new_config["progress_throttle"] = 0.1
            return new_config
        elif old_config.get("version", 3) == 3:
            new_config = old_config.copy()
            new_config["version"] = 4
            new_config["enable_sounds"] = True
            new_config["show_notifications"] = True
            new_config["auto_load_last_files"] = True
            new_config["ocr_enabled"] = False
            new_config["tesseract_path"] = self._detect_tesseract_path()
            new_config["max_history_items"] = 50
            return new_config
        return old_config
    
    def _detect_libreoffice_path(self) -> str:
        """Tries to detect LibreOffice path automatically"""
        try:
            if platform.system() == "Windows":
                # Common Windows paths
                paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                for path in paths:
                    if os.path.exists(path):
                        return path
            else:
                # Linux/Mac - assumes it's in PATH
                result = subprocess.run(["which", "soffice"], capture_output=True, text=True)
                if result.returncode == 0:
                    return result.stdout.strip()
        except Exception:
            pass
        return ""
    
    def _detect_tesseract_path(self) -> str:
        """Tries to detect Tesseract OCR path automatically"""
        try:
            if platform.system() == "Windows":
                # Common Windows paths
                paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
                ]
                for path in paths:
                    if os.path.exists(path):
                        return path
            else:
                # Linux/Mac - assumes it's in PATH
                result = subprocess.run(["which", "tesseract"], capture_output=True, text=True)
                if result.returncode == 0:
                    return result.stdout.strip()
        except Exception:
            pass
        return ""
    
    def save_config(self):
        """Saves configuration to file with robust error handling"""
        try:
            temp_file = f"{self.CONFIG_FILE}.tmp"
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
            
            if os.path.exists(temp_file):
                if os.path.exists(self.CONFIG_FILE):
                    os.replace(temp_file, self.CONFIG_FILE)
                else:
                    os.rename(temp_file, self.CONFIG_FILE)
        except Exception as e:
            logger.error(f"Error saving config: {str(e)}")
            logger.debug(traceback.format_exc())
    
    def update_recent_file(self, file_type: str, path: Union[str, List[str]]):
        """Updates recent path for a file type"""
        if file_type in self.config["recent_files"]:
            if isinstance(path, list):
                current = self.config["recent_files"][file_type]
                if not isinstance(current, list):
                    current = []
                
                updated = current + path
                updated = list(dict.fromkeys(updated))
                self.config["recent_files"][file_type] = updated[-5:]
            else:
                self.config["recent_files"][file_type] = path
            self.save_config()
    
    def save_column_mapping(self, template_path: str, mapping: Dict[str, str]):
        """Saves column mapping for a specific template"""
        self.config["column_mappings"][template_path] = mapping
        self.save_config()
    
    def load_column_mapping(self, template_path: str) -> Optional[Dict[str, str]]:
        """Loads column mapping for a specific template"""
        return self.config["column_mappings"].get(template_path)

class DocumentProcessor:
    """Base class for document processing with performance optimizations"""
    
    def __init__(self, max_workers: int = 4):
        self.running = False
        self.stop_requested = False
        self.progress_queue = queue.Queue()
        self.log_queue = queue.Queue(maxsize=1000)  # Limits log size in memory
        self.error_queue = queue.Queue()
        self.gui_update_queue = queue.Queue()
        self.max_workers = max_workers
        self.processed_count = 0
        self.error_count = 0
        self.start_time = None
        self.end_time = None
        self.last_progress_update = 0
        self.progress_throttle = 0.1  # Minimum time between progress updates
        self.operation_type = "Generic Operation"

    def check_system_resources(self):
        """Verifica recursos antes de processar"""
        mem = psutil.virtual_memory()
        if mem.available < 2 * 1024 * 1024 * 1024:  # 2GB livres
            self.log_message("⚠️ Memória insuficiente - reduza o lote", "warning")
            return False
        
        cpu_load = psutil.cpu_percent(interval=1)
        if cpu_load > 80:  # 80% de uso
            self.log_message("⚠️ CPU sobrecarregada - aguardando...", "warning")
            time.sleep(5)
            return self.check_system_resources()
        
        return True
    
    def process_large_batch(self, file_list, batch_size=100):
        for i in range(0, len(file_list), batch_size):
            batch = file_list[i:i + batch_size]
            if not self.check_system_resources():
                time.sleep(10)
            
            # Processa o lote atual
            self.process_batch(batch)
            
            # Limpa recursos
            gc.collect()
            time.sleep(2)

    def process_document(self, *args, **kwargs):
        raise NotImplementedError
    
    def safe_process(self, func, *args, **kwargs):
        try:
            logger.debug(f"Executing {func.__name__} with args: {args}, kwargs: {kwargs}")
            
            # Add timeout for I/O operations
            if func.__name__ in ['read_excel', 'save', 'convert']:
                with ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(func, *args, **kwargs)
                    result = future.result(timeout=30)  # 30 second timeout
            else:
                result = func(*args, **kwargs)
                
            logger.debug(f"Result from {func.__name__}: {result}")
            return result
        except Exception as e:
            error_msg = f"ERROR in {func.__name__}: {str(e)}"
            logger.error(error_msg)
            logger.debug(traceback.format_exc())
            self.log_message(error_msg, level="error")
            self.error_queue.put((args, kwargs, str(e), traceback.format_exc()))
            return None
    
    def log_message(self, message: str, level: str = "info"):
        """Adds message to log with size limit"""
        if self.log_queue.full():
            self.log_queue.get()  # Removes oldest message if queue is full
        self.log_queue.put((message, level))
    
    def update_progress(self, value: int = 1, max_value: int = None):
        """Updates progress with throttling to avoid overload"""
        now = time.time()
        if now - self.last_progress_update >= self.progress_throttle:
            self.progress_queue.put((value, max_value))
            self.last_progress_update = now
    
    def safe_gui_update(self, widget, method, *args, **kwargs):
        """Puts GUI updates in a safe queue for execution in main thread"""
        self.gui_update_queue.put((widget, method, args, kwargs))
    
    def check_stop_requested(self) -> bool:
        return self.stop_requested
    
    def stop(self):
        self.stop_requested = True
        self.running = False
        self.log_message("Processing stopped by user", "warning")
    
    def _generate_execution_report(self) -> Dict:
        """Generates a detailed execution report"""
        duration = (self.end_time - self.start_time).total_seconds() if self.start_time and self.end_time else 0
        
        return {
            "operation_type": self.operation_type,
            "start_time": self.start_time.isoformat() if self.start_time else None,
            "end_time": self.end_time.isoformat() if self.end_time else None,
            "duration": duration,
            "processed_count": self.processed_count,
            "error_count": self.error_count,
            "success_rate": (self.processed_count / (self.processed_count + self.error_count)) * 100 
                           if (self.processed_count + self.error_count) > 0 else 0,
            "status": "completed" if not self.stop_requested else "stopped"
        }
    
    def _play_completion_sound(self):
        """Plays a sound notification when processing completes"""
        try:
            if platform.system() == "Windows":
                winsound.MessageBeep()
            else:
                # For Mac/Linux - simple beep
                print("\a")
        except Exception as e:
            logger.error(f"Error playing sound: {str(e)}")

class BatchDeclarationGenerator(DocumentProcessor):
    """Manages batch declaration generation with optimizations"""
    
    def __init__(self, max_workers: int = 4, use_libreoffice: bool = False, 
                libreoffice_path: str = ""):
        super().__init__(max_workers)
        self.placeholders = []
        self.column_mapping = {}
        self.operation_type = "Declaration Generation"
        self.document_manager = DocumentManager(use_libreoffice, libreoffice_path)
        self.save_in_subfolders = False  # Ativar via interface
        self.subfolder_column = "Nome"  # Coluna para criar subpastas
        self.custom_filename_pattern = "{{Nome}}_{{TipoDoc}}"  # Padrão configurável


    
    def detect_placeholders(self, template_path: str) -> List[str]:
        if not os.path.exists(template_path):
            self.log_message(f"Template file not found: {template_path}", "error")
            return []
        
        placeholders = set()
        
        def extract_from_paragraph(paragraph):
            text = paragraph.text
            start = text.find("{{")
            while start != -1:
                end = text.find("}}", start)
                if end != -1:
                    placeholder = text[start:end+2]
                    placeholders.add(placeholder)
                    start = text.find("{{", end)
                else:
                    self.log_message(f"Unclosed placeholder: {text[start:start+20]}...", "warning")
                    break
        
        try:
            doc = self.safe_process(DocxDocument, template_path)
            if not doc:
                return []
            
            for para in doc.paragraphs:
                self.safe_process(extract_from_paragraph, para)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.safe_process(extract_from_paragraph, para)
            
            self.placeholders = sorted(list(placeholders))
            
            # Validate placeholders for common issues
            invalid_placeholders = []
            for ph in self.placeholders:
                if not re.match(r"^\{\{[A-Za-z0-9_\- ]+\}\}$", ph):
                    invalid_placeholders.append(ph)
            
            if invalid_placeholders:
                self.log_message(f"Warning: Found potentially invalid placeholders: {', '.join(invalid_placeholders)}", "warning")
            
            return self.placeholders
            
        except Exception as e:
            self.log_message(f"Error analyzing template: {str(e)}", "error")
            logger.exception("Error detecting placeholders")
            return []
    
    def validate_columns(self, excel_path: str, template_path: str, config: EnhancedConfigManager) -> Tuple[bool, Dict[str, str]]:
        if not os.path.exists(excel_path):
            self.log_message(f"Excel file not found: {excel_path}", "error")
            return False, {}
        
        try:
            # Try to load saved mapping
            saved_mapping = config.load_column_mapping(template_path)
            
            df = self.safe_process(pd.read_excel, excel_path)
            if df is None or df.empty:
                self.log_message("Empty or invalid Excel sheet", "error")
                return False, {}
            
            column_mapping = {}
            missing = []
            warnings = []
            suggestions = {}
            
            for ph in self.placeholders:
                col_name = ph[2:-2].strip()
                
                # Check saved mapping first
                if saved_mapping and ph in saved_mapping:
                    mapped_col = saved_mapping[ph]
                    if mapped_col in df.columns:
                        column_mapping[ph] = mapped_col
                        continue
                    else:
                        warnings.append(f"Mapped column '{mapped_col}' not found in Excel. Looking for alternatives.")
                
                if col_name in df.columns:
                    column_mapping[ph] = col_name
                else:
                    # Enhanced column matching with multiple strategies
                    matches = self._find_best_column_match(col_name, df.columns)
                    if matches:
                        suggestions[col_name] = matches
                        column_mapping[ph] = matches[0]
                        warnings.append(f"Column '{col_name}' not found. Using '{matches[0]}'")
                    else:
                        missing.append(col_name)
            
            self.column_mapping = column_mapping
            
            result = {
                "mapping": column_mapping,
                "missing": missing,
                "warnings": warnings,
                "suggestions": suggestions,
                "columns": list(df.columns),
                "saved_mapping_used": saved_mapping is not None
            }
            
            return len(missing) == 0, result
            
        except Exception as e:
            self.log_message(f"Error validating columns: {str(e)}", "error")
            logger.exception("Error in column validation")
            return False, {}
    
    def _find_best_column_match(self, target: str, columns: List[str]) -> List[str]:
        """Finds the best matching column using multiple strategies"""
        target_lower = target.lower().strip()
        columns_lower = [c.lower().strip() for c in columns]
        
        # 1. Exact match (case insensitive)
        if target_lower in columns_lower:
            return [columns[columns_lower.index(target_lower)]]
        
        # 2. Close matches using difflib
        matches = get_close_matches(target_lower, columns_lower, n=3, cutoff=0.6)
        if matches:
            return [columns[columns_lower.index(m)] for m in matches]
        
        # 3. Partial matches (contains)
        partial_matches = [c for c in columns if target_lower in c.lower() or c.lower() in target_lower]
        if partial_matches:
            return partial_matches
        
        # 4. Split words matching
        target_words = set(target_lower.split())
        word_matches = []
        for c in columns:
            col_words = set(c.lower().split())
            if target_words & col_words:  # Any common words
                word_matches.append(c)
        
        if word_matches:
            return word_matches
        
        return []
    
    def generate_declarations(self, excel_path: str, template_path: str, 
                        output_folder: str, create_zip: bool = False, 
                        config: EnhancedConfigManager = None,
                        nome_saida_template: str = "{{Nome}}_{{TipoDoc}}"):  

        self.running = True
        self.stop_requested = False
        self.processed_count = 0
        self.error_count = 0
        self.start_time = datetime.now()
        
        try:
            if not all(map(os.path.exists, [excel_path, template_path])):
                raise FileNotFoundError("One or more input files not found")
            
            if not os.path.isdir(output_folder):
                os.makedirs(output_folder, exist_ok=True)
                self.log_message(f"Created output folder: {output_folder}")
            
            self.log_message("Reading Excel spreadsheet...")
            df = self.safe_process(pd.read_excel, excel_path)
            if df is None:
                raise ValueError("Error reading Excel file")
            
            total = len(df)
            if total == 0:
                raise ValueError("Excel sheet contains no data")
            
            self.update_progress(0, total)
            
            chunk_size = 50  # define o tamanho de cada lote

            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                for start in range(0, len(df), chunk_size):
                    chunk = df.iloc[start:start+chunk_size]
                    futures = []

                    for index, row in chunk.iterrows():
                        if self.check_stop_requested():
                            break
                        
                        futures.append(
                            executor.submit(
                                self._generate_single_declaration,
                                index, row, template_path, output_folder, nome_saida_template
                            )
                        )


                    for future in as_completed(futures):
                        if self.check_stop_requested():
                            break

                        try:
                            result = future.result()
                            if result:
                                self.processed_count += 1
                                self.log_message(f"Generated: {os.path.basename(result)}")
                        except Exception as e:
                            self.error_count += 1
                            self.log_message(f"Error processing document: {str(e)}", "error")

            
            if create_zip and not self.check_stop_requested():
                self._create_zip_file(output_folder)
            
            # Save column mapping if config was provided
            if config and self.column_mapping:
                config.save_column_mapping(template_path, self.column_mapping)
            
            self.end_time = datetime.now()
            report = self._generate_final_report(total)
            
            # # Generate CSV report
            # csv_report_path = os.path.join(output_folder, "processing_report.csv")
            # self._generate_csv_report(csv_report_path, total)
            
            return self.processed_count, self.error_count
            
        except Exception as e:
            self.log_message(f"❌ CRITICAL ERROR: {str(e)}", "error")
            logger.exception("Error during declaration generation")
            return 0, 1
            
        finally:
            self.running = False
            if config and config.config.get("enable_sounds", True):
                self._play_completion_sound()
                
    def _generate_single_declaration(self, index: int, row: pd.Series, 
                                    template_path: str, output_folder: str,
                                    nome_saida_template: str) -> Optional[str]:
        if self.check_stop_requested():
            self.log_message(f"⏹️ Cancelado antes de processar a linha {index+1}")
            return None

        try:
            nome_arquivo = nome_saida_template
            for col in row.index:
                placeholder = f"{{{{{col}}}}}"
                value = str(row[col]).strip()
                nome_arquivo = nome_arquivo.replace(placeholder, value)
            
            nome_arquivo = re.sub(r'[\\/*?:"<>|]', "", nome_arquivo)
            nome_arquivo = nome_arquivo.replace(" ", "_") + ".docx"

            if self.save_in_subfolders:
                subfolder = str(row.get('Nome', '')).strip()
                subfolder = re.sub(r'[\\/*?:"<>|]', "", subfolder)
                if subfolder:
                    output_folder = os.path.join(output_folder, subfolder)
                    os.makedirs(output_folder, exist_ok=True)

            if self.check_stop_requested():
                self.log_message(f"⏹️ Cancelado antes de abrir o template - linha {index+1}")
                return None

            doc = self.document_manager.open_document(template_path)
            if doc is None:
                raise ValueError("Falha ao abrir o template")

            if self.check_stop_requested():
                self.log_message(f"⏹️ Cancelado antes de aplicar substituições - linha {index+1}")
                return None

            replacements = {f"{{{{{col}}}}}": str(row[col]) for col in row.index}
            self._apply_replacements_word(doc, replacements)

            if self.check_stop_requested():
                self.log_message(f"⏹️ Cancelado antes de salvar - linha {index+1}")
                return None

            output_path = os.path.join(output_folder, nome_arquivo)
            self.document_manager.save_document(doc, output_path)
            doc.close()

            return output_path

        except Exception as e:
            self.log_message(f"ERRO linha {index+1}: {str(e)}", "error")
            return None


    def _sanitize_filename(self, filename: str) -> str:
        """Remove caracteres inválidos e formata o nome do arquivo"""
        # Substitui espaços por underscores se preferir
        filename = filename.replace(' ', '_')
        
        # Remove caracteres inválidos
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '')
        
        # Remove múltiplos underscores consecutivos
        filename = re.sub(r'_+', '_', filename)
        
        return filename.strip('_')

    
    from docx.text.run import Run

    def _apply_replacements_word(self, doc, replacements: Dict[str, str]):
        def substituir_em_runs(paragraph):
            for run in paragraph.runs:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)

        # Substituir nos parágrafos
        for paragraph in doc.paragraphs:
            substituir_em_runs(paragraph)

        # Substituir nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        substituir_em_runs(paragraph)



    def _apply_replacements_libreoffice(self, doc: LibreOfficeDocument, replacements: Dict[str, str]):
        # Implementação melhorada para LibreOffice
        try:
            # 1. Converter para ODT (formato aberto)
            odt_path = os.path.join(doc.temp_dir, "temp_replace.odt")
            doc.save_as(odt_path)
            
            # 2. Extrair conteúdo XML
            with zipfile.ZipFile(odt_path, 'r') as z:
                content = z.read('content.xml').decode('utf-8')
            
            # 3. Aplicar substituições
            for ph, value in replacements.items():
                content = content.replace(ph, value)
            
            # 4. Recriar arquivo ODT
            with zipfile.ZipFile(odt_path, 'w') as z:
                for f in z.filelist:
                    if f.filename != 'content.xml':
                        z.writestr(f, z.read(f))
                z.writestr('content.xml', content.encode('utf-8'))
            
            # 5. Converter de volta para DOCX
            output_path = os.path.join(doc.temp_dir, "final.docx")
            cmd = [
                self.libreoffice_path,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(output_path),
                odt_path
            ]
            subprocess.run(cmd, check=True)
            
            return True
        except Exception as e:
            logger.error(f"LibreOffice replacement error: {str(e)}")
            return False
    
    # def _prepare_replacements(self, row: pd.Series) -> Dict[str, str]:
    #     replacements = {
    #         "{{NOME}}": str(row.get("Nome", "")).strip(),
    #         "{{CPF}}": str(row.get("CPF", "")).strip(),
    #         "{{MATRICULA}}": str(row.get("Matricula", "")).strip(),
    #         "{{EMAIL}}": str(row.get("Email", "")).strip(),
    #         "{{TIPO_DOC}}": str(row.get("TipoDoc", "")).strip(),
    #         "{{TITULO}}": str(row.get("Titulo", "")).strip(),
    #         "{{DATA_DEFESA}}": str(row.get("DataDefesa", "")).strip(),
    #         "{{LOCAL}}": str(row.get("Local", "")).strip(),
    #         "{{HORARIO}}": str(row.get("Horario", "")).strip(),
    #         "{{NOME_ORIENTADOR}}": str(row.get("NomeOrientador", "")).strip(),
    #         "{{NOME_COORIENTADOR}}": str(row.get("NomeCoorientador", "")).strip(),
    #         "{{NOME_EXAMINADOR1}}": str(row.get("NomeExaminador1", "")).strip(),
    #         "{{NOME_EXAMINADOR2}}": str(row.get("NomeExaminador2", "")).strip(),
           

    #     }
        
    #     for col in row.index:
    #         replacements[f"{{{{{col.upper()}}}}}"] = str(row[col]) if pd.notna(row[col]) else ""

        
    #     return replacements
    

    def _prepare_replacements(self, row: pd.Series) -> Dict[str, str]:
        replacements = {}
        
        # Adiciona todos os campos da linha
        for col in row.index:
            value = str(row[col]) if pd.notna(row[col]) else ""
            # Adiciona variações do placeholder
            replacements[f"{{{{{col}}}}}"] = value                 # Ex: {{Local}}
            replacements[f"{{{{{col.upper()}}}}}"] = value         # Ex: {{LOCAL}}
            replacements[f"{{{{{col.lower()}}}}}"] = value         # Ex: {{local}}
        
        # Adiciona campos padrão adicionais
        now = datetime.now()
        replacements.update({
            # "{{DATA}}": now.strftime("%d/%m/%Y"),
            # "{{DIA}}": str(now.day),
            # "{{MES}}": str(now.month),
            # "{{ANO}}": str(now.year),
            # "{{HORA}}": now.strftime("%H:%M"),
        })
        
        return replacements



    def _apply_replacements(self, doc: DocxDocument, replacements: Dict[str, str]):
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)
    
    def _generate_safe_filename(self, name: str) -> str:
        keep_chars = (' ', '.', '_', '-')
        safe_name = "".join(c for c in str(name) if c.isalnum() or c in keep_chars).strip()
        return safe_name[:100]
    
    def _create_zip_file(self, output_folder: str):
        self.log_message("Creating ZIP archive...")
        zip_path = os.path.join(output_folder, "declarações.zip")
        
        try:
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, _, files in os.walk(output_folder):
                    for file in files:
                        if file.endswith('.docx') and not file.startswith('~$'):
                            file_path = os.path.join(root, file)
                            zipf.write(file_path, os.path.basename(file_path))
            
            self.log_message(f"ZIP file created: {os.path.basename(zip_path)}", "success")
        except Exception as e:
            self.log_message(f"Error creating ZIP: {str(e)}", "error")
            logger.exception("Error creating ZIP file")
    
    def _generate_final_report(self, total: int):
        self.log_message("\n=== FINAL REPORT ===")
        self.log_message(f"📊 Total documents: {total}")
        self.log_message(f"✅ Successes: {self.processed_count}", "success")
        
        if self.error_count > 0:
            self.log_message(f"⚠️ Errors: {self.error_count}", "warning")
        
        if self.start_time and self.end_time:
            duration = self.end_time - self.start_time
            self.log_message(f"⏱️ Total time: {duration}")
        
        if self.stop_requested:
            self.log_message("🛑 Processing stopped by user", "warning")
        
        return self._generate_execution_report()
    
    def _generate_pdf_report(self, output_folder: str):
        """Generates a PDF report with processing results"""
        report_path = os.path.join(output_folder, "processing_report.pdf")
        
        try:
            c = canvas.Canvas(report_path, pagesize=letter)
            width, height = letter
            
            # Header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, height - 72, "Processing Report")
            
            # Basic info
            c.setFont("Helvetica", 12)
            y_position = height - 100
            c.drawString(72, y_position, f"Total documents: {self.processed_count + self.error_count}")
            y_position -= 20
            c.drawString(72, y_position, f"Successes: {self.processed_count}")
            y_position -= 20
            c.drawString(72, y_position, f"Errors: {self.error_count}")
            y_position -= 20
            
            if self.start_time and self.end_time:
                c.drawString(72, y_position, f"Total time: {self.end_time - self.start_time}")
                y_position -= 30
            
            # Error details (if any)
            if self.error_count > 0:
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, y_position, "Error Details:")
                y_position -= 20
                c.setFont("Helvetica", 10)
                
                c.drawString(72, y_position, "Check the log file for complete details.")
                y_position -= 20
            
            c.save()
            self.log_message(f"PDF report generated: {report_path}", "success")
            return report_path
        except Exception as e:
            self.log_message(f"Error generating PDF report: {str(e)}", "error")
            return None
    
    def _generate_csv_report(self, output_path: str, total: int):
        """Generates a CSV report with processing results"""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["Metric", "Value"])
                writer.writerow(["Total Documents", total])
                writer.writerow(["Successes", self.processed_count])
                writer.writerow(["Errors", self.error_count])
                writer.writerow(["Success Rate", f"{(self.processed_count / total * 100):.2f}%"])
                
                if self.start_time and self.end_time:
                    duration = (self.end_time - self.start_time).total_seconds()
                    writer.writerow(["Duration (seconds)", f"{duration:.2f}"])
                
                writer.writerow(["Status", "Completed" if not self.stop_requested else "Stopped"])
            
            self.log_message(f"CSV report generated: {output_path}", "success")
            return output_path
        except Exception as e:
            self.log_message(f"Error generating CSV report: {str(e)}", "error")
            return None
        
    def _generate_custom_filename(self, row: pd.Series, pattern: str) -> str:
        """Gera o nome do arquivo baseado no padrão fornecido pelo usuário"""
        filename = pattern
        for col in row.index:
            placeholder = f"{{{{{col}}}}}"
            value = str(row[col]).strip()
            filename = filename.replace(placeholder, value)
        
        keep_chars = (' ', '.', '_', '-', '(', ')')
        safe_name = "".join(c for c in filename if c.isalnum() or c in keep_chars).strip()
        return safe_name[:100]  # Limita tamanho

    def _generate_output_folder(self, base_output_folder: str, row: pd.Series, folder_column: str) -> str:
        """Cria subpasta com base no valor de uma coluna"""
        folder_name = str(row.get(folder_column, "SemNome")).strip()
        folder_name = "".join(c for c in folder_name if c.isalnum() or c in (' ', '_', '-')).strip()
        target_folder = os.path.join(base_output_folder, folder_name)
        os.makedirs(target_folder, exist_ok=True)
        return target_folder

class DocumentConverter(DocumentProcessor):
    """Batch document converter with optimizations"""
    
    def __init__(self, max_workers: int = 4, use_libreoffice: bool = False, 
                libreoffice_path: str = "", ocr_enabled: bool = False, 
                tesseract_path: str = ""):
        super().__init__(max_workers)
        self.use_libreoffice = use_libreoffice
        self.libreoffice_path = libreoffice_path
        self.ocr_enabled = ocr_enabled
        self.tesseract_path = tesseract_path

    def _is_valid_docx(self, file_path: str) -> bool:
        """Verifica se é um arquivo .docx válido com estrutura interna"""
        if not file_path.lower().endswith('.docx'):
            return False
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                return 'word/document.xml' in z.namelist()
        except:
            return False

    
    def convert_docx_to_pdf(self, docx_paths: List[str], output_folder: str):
        self.operation_type = "DOCX to PDF Conversion"
        self.running = True
        self.stop_requested = False
        self.processed_count = 0
        self.error_count = 0
        self.start_time = datetime.now()
        
        try:
            valid_paths = [
                p for p in docx_paths 
                if os.path.exists(p)
                and p.lower().endswith('.docx')
                and self._is_valid_docx(p)
            ]

            ignored = [p for p in docx_paths if p not in valid_paths]
            if ignored:
                self.log_message(f"Ignored {len(ignored)} files (invalid format or not real .docx)", "warning")


            if not valid_paths:
                raise FileNotFoundError("No valid DOCX files found")
            
            if len(valid_paths) != len(docx_paths):
                self.log_message(f"Ignored {len(docx_paths)-len(valid_paths)} invalid files", "warning")
            
            if not os.path.isdir(output_folder):
                os.makedirs(output_folder, exist_ok=True)
                self.log_message(f"Created output folder: {output_folder}")
            
            total = len(valid_paths)
            self.update_progress(0, total)
            
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                futures = []
                
                for docx_path in valid_paths:
                    if self.check_stop_requested():
                        break
                    
                    futures.append(
                        executor.submit(
                            self._convert_single_docx_to_pdf,
                            docx_path, output_folder
                        )
                    )
                
                for future in as_completed(futures):
                    if self.check_stop_requested():
                        break
                    
                    try:
                        result = future.result()
                        if result:
                            self.processed_count += 1
                            self.log_message(f"Converted: {os.path.basename(result)}")
                    except Exception as e:
                        self.error_count += 1
                        self.log_message(f"Conversion error: {str(e)}", "error")
            
            self.end_time = datetime.now()
            report = self._generate_conversion_report("DOCX to PDF", total)
            
            # Generate CSV report
            # csv_report_path = os.path.join(output_folder, "conversion_report.csv")
            # self._generate_csv_report(csv_report_path, "DOCX to PDF", total)
            
            return self.processed_count, self.error_count
            
        except Exception as e:
            self.log_message(f"❌ CRITICAL ERROR: {str(e)}", "error")
            logger.exception("Error during DOCX to PDF conversion")
            return 0, 1
            
        finally:
            self.running = False
        

    
    def _convert_single_docx_to_pdf(self, docx_path: str, output_folder: str) -> Optional[str]:
        """Converte um único arquivo DOCX para PDF com tratamento robusto de erros"""
        try:
            # Verifica se o arquivo DOCX existe
            if not os.path.exists(docx_path):
                self.log_message(f"Arquivo não encontrado: {docx_path}", "error")
                return None

            # Cria a pasta de saída se não existir
            os.makedirs(output_folder, exist_ok=True)

            # Nome do arquivo de saída
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            output_path = os.path.join(output_folder, f"{base_name}.pdf")

            # Verifica se o arquivo de saída já existe e cria um nome único
            counter = 1
            while os.path.exists(output_path):
                output_path = os.path.join(output_folder, f"{base_name}_{counter}.pdf")
                counter += 1

            # Verifica se o LibreOffice está configurado
            if not self.libreoffice_path or not os.path.exists(self.libreoffice_path):
                self.log_message("Caminho do LibreOffice não configurado ou inválido", "error")
                return None

            # Prepara o comando - usa shlex.quote para lidar com espaços no caminho
            cmd = [
                self.libreoffice_path,
                '--headless',
                '--convert-to', 'pdf:writer_pdf_Export',
                '--outdir', output_folder,
                docx_path
            ]

            self.log_message(f"Convertendo: {os.path.basename(docx_path)}...", "info")

            # Executa o comando com timeout
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,  # 2 minutos de timeout
                text=True,
                encoding='utf-8',
                errors='ignore'
            )

            # Verifica o resultado
            if result.returncode != 0:
                error_msg = result.stderr or "Erro desconhecido"
                self.log_message(f"Erro na conversão: {error_msg}", "error")
                return None

            # Verifica se o arquivo foi criado
            if not os.path.exists(output_path):
                self.log_message(f"Arquivo PDF não foi gerado: {output_path}", "error")
                return None

            self.log_message(f"Conversão bem-sucedida: {os.path.basename(output_path)}", "success")
            return output_path

        except subprocess.TimeoutExpired:
            self.log_message(f"Timeout ao converter {os.path.basename(docx_path)}", "error")
            return None
        except Exception as e:
            self.log_message(f"Erro inesperado ao converter {os.path.basename(docx_path)}: {str(e)}", "error")
            logger.exception("Erro na conversão DOCX para PDF")
            return None

    
    def _convert_with_libreoffice(self, docx_path: str, output_path: str):
        """Converts DOCX to PDF using LibreOffice"""
        try:
            cmd = [
                self.libreoffice_path,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(output_path),
                docx_path
            ]
            
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60  # 60 second timeout
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice returned code {result.returncode}: {result.stderr.decode()}")
            
        except Exception as e:
            raise RuntimeError(f"LibreOffice conversion failed: {str(e)}")
    
    def convert_pdf_to_docx(self, pdf_paths: List[str], output_folder: str):
        self.operation_type = "PDF to DOCX Conversion"
        self.running = True
        self.stop_requested = False
        self.processed_count = 0
        self.error_count = 0
        self.start_time = datetime.now()
        
        try:
            valid_paths = [p for p in pdf_paths if os.path.exists(p)]
            if not valid_paths:
                raise FileNotFoundError("No valid PDF files found")
            
            if len(valid_paths) != len(pdf_paths):
                self.log_message(f"Ignored {len(pdf_paths)-len(valid_paths)} invalid files", "warning")
            
            if not os.path.isdir(output_folder):
                os.makedirs(output_folder, exist_ok=True)
                self.log_message(f"Created output folder: {output_folder}")
            
            total = len(valid_paths)
            self.update_progress(0, total)
            
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                futures = []
                
                for pdf_path in valid_paths:
                    if self.check_stop_requested():
                        break
                    
                    futures.append(
                        executor.submit(
                            self._convert_single_pdf_to_docx,
                            pdf_path, output_folder
                        )
                    )
                
                for future in as_completed(futures):
                    if self.check_stop_requested():
                        break
                    
                    try:
                        result = future.result()
                        if result:
                            self.processed_count += 1
                            self.log_message(f"Converted: {os.path.basename(result)}")
                    except Exception as e:
                        self.error_count += 1
                        self.log_message(f"Conversion error: {str(e)}", "error")
            
            self.end_time = datetime.now()
            report = self._generate_conversion_report("PDF to DOCX", total)
            
            # Generate CSV report
            # csv_report_path = os.path.join(output_folder, "conversion_report.csv")
            # self._generate_csv_report(csv_report_path, "PDF to DOCX", total)
            
            return self.processed_count, self.error_count
            
        except Exception as e:
            self.log_message(f"❌ CRITICAL ERROR: {str(e)}", "error")
            logger.exception("Error during PDF to DOCX conversion")
            return 0, 1
            
        finally:
            self.running = False
    
    def _convert_single_pdf_to_docx(self, pdf_path: str, output_folder: str) -> Optional[str]:
        try:
            output_path = os.path.join(
                output_folder, 
                os.path.splitext(os.path.basename(pdf_path))[0] + ".docx"
            )
            
            if os.path.exists(output_path):
                output_path = self._get_unique_filename(output_path)
            
            self.log_message(f"Converting: {os.path.basename(pdf_path)}")
            
            if self.ocr_enabled and self._is_scanned_pdf(pdf_path):
                self.log_message(f"Using OCR for scanned PDF: {os.path.basename(pdf_path)}", "info")
                return self._convert_with_ocr(pdf_path, output_path)
            else:
                cv = Converter(pdf_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            
            self.update_progress(1)
            return output_path
            
        except Exception as e:
            self.log_message(f"Error converting {os.path.basename(pdf_path)}: {str(e)}", "error")
            logger.error(f"Error converting {pdf_path}")
            logger.debug(traceback.format_exc())
            self.error_count += 1
            return None
    
    def _is_scanned_pdf(self, pdf_path: str) -> bool:
        """Checks if a PDF is likely scanned (image-based)"""
        try:
            # Simple check - if the first page has very little text, assume it's scanned
            cv = Converter(pdf_path)
            text = cv.extract_text(0, 1)  # Extract text from first page
            cv.close()
            
            return len(text.strip()) < 50  # Arbitrary threshold
        except Exception:
            return False
    
    def _convert_with_ocr(self, pdf_path: str, output_path: str) -> str:
        """Converts a scanned PDF to DOCX using OCR"""
        if not self.tesseract_path or not os.path.exists(self.tesseract_path):
            raise RuntimeError("Tesseract OCR path not configured or invalid")
        
        try:
            # First convert PDF to images
            images = self._pdf_to_images(pdf_path)
            
            # Then create a new Word document with OCR text
            doc = DocxDocument()
            
            for img_path in images:
                # Perform OCR on each image
                text = self._perform_ocr(img_path)
                
                # Add text to document
                if text.strip():
                    doc.add_paragraph(text)
                
                # Add the original image
                doc.add_picture(img_path)
                
                # Clean up temporary image file
                os.unlink(img_path)
            
            doc.save(output_path)
            return output_path
        except Exception as e:
            # Clean up any remaining temp files
            if 'images' in locals():
                for img_path in images:
                    if os.path.exists(img_path):
                        try:
                            os.unlink(img_path)
                        except:
                            pass
            raise
    
    def _pdf_to_images(self, pdf_path: str) -> List[str]:
        """Converts PDF pages to temporary image files"""
        try:
            images = []
            temp_dir = tempfile.mkdtemp()
            
            # Use pdf2image or similar library here
            # For simplicity, we'll just generate placeholder code
            # In a real implementation, you would use something like:
            # from pdf2image import convert_from_path
            # images = convert_from_path(pdf_path, output_folder=temp_dir, fmt='png')
            
            # Placeholder - replace with actual implementation
            raise NotImplementedError("PDF to image conversion not implemented")
            
            return [img.filename for img in images]
        except Exception as e:
            # Clean up temp dir
            if os.path.exists(temp_dir):
                try:
                    for f in os.listdir(temp_dir):
                        os.unlink(os.path.join(temp_dir, f))
                    os.rmdir(temp_dir)
                except:
                    pass
            raise
    
    def _perform_ocr(self, image_path: str) -> str:
        """Performs OCR on an image file"""
        try:
            # Configure Tesseract path if needed
            if self.tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
            
            # Perform OCR
            text = pytesseract.image_to_string(Image.open(image_path))
            return text
        except Exception as e:
            raise RuntimeError(f"OCR failed: {str(e)}")
    
    def _get_unique_filename(self, path: str) -> str:
        base, ext = os.path.splitext(path)
        counter = 1
        new_path = path
        
        while os.path.exists(new_path):
            new_path = f"{base}_{counter}{ext}"
            counter += 1
        
        return new_path
    
    def _generate_conversion_report(self, operation: str, total: int):
        self.log_message("\n=== CONVERSION REPORT ===")
        self.log_message(f"Operation: {operation}")
        self.log_message(f"📊 Total files: {total}")
        self.log_message(f"✅ Successes: {self.processed_count}", "success")
        
        if self.error_count > 0:
            self.log_message(f"⚠️ Errors: {self.error_count}", "warning")
        
        if self.start_time and self.end_time:
            duration = self.end_time - self.start_time
            self.log_message(f"⏱️ Total time: {duration}")
        
        if self.stop_requested:
            self.log_message("🛑 Processing stopped by user", "warning")
        
        return self._generate_execution_report()
    
    def _generate_pdf_report(self, output_folder: str, operation: str):
        """Generates a PDF report with conversion results"""
        report_path = os.path.join(output_folder, f"conversion_report_{operation.replace(' ', '_')}.pdf")
        
        try:
            c = canvas.Canvas(report_path, pagesize=letter)
            width, height = letter
            
            # Header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, height - 72, f"Conversion Report - {operation}")
            
            # Basic info
            c.setFont("Helvetica", 12)
            y_position = height - 100
            c.drawString(72, y_position, f"Total files: {self.processed_count + self.error_count}")
            y_position -= 20
            c.drawString(72, y_position, f"Successful conversions: {self.processed_count}")
            y_position -= 20
            c.drawString(72, y_position, f"Errors: {self.error_count}")
            y_position -= 20
            
            if self.start_time and self.end_time:
                c.drawString(72, y_position, f"Total time: {self.end_time - self.start_time}")
                y_position -= 30
            
            # Error details (if any)
            if self.error_count > 0:
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, y_position, "Error Details:")
                y_position -= 20
                c.setFont("Helvetica", 10)
                
                c.drawString(72, y_position, "Check the log file for complete details.")
                y_position -= 20
            
            c.save()
            self.log_message(f"PDF report generated: {report_path}", "success")
            return report_path
        except Exception as e:
            self.log_message(f"Error generating PDF report: {str(e)}", "error")
            return None
    
    def _generate_csv_report(self, output_path: str, operation: str, total: int):
        """Generates a CSV report with conversion results"""
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow(["Metric", "Value"])
                writer.writerow(["Operation", operation])
                writer.writerow(["Total Files", total])
                writer.writerow(["Successes", self.processed_count])
                writer.writerow(["Errors", self.error_count])
                writer.writerow(["Success Rate", f"{(self.processed_count / total * 100):.2f}%"])
                
                if self.start_time and self.end_time:
                    duration = (self.end_time - self.start_time).total_seconds()
                    writer.writerow(["Duration (seconds)", f"{duration:.2f}"])
                
                writer.writerow(["Status", "Completed" if not self.stop_requested else "Stopped"])
            
            self.log_message(f"CSV report generated: {output_path}", "success")
            return output_path
        except Exception as e:
            self.log_message(f"Error generating CSV report: {str(e)}", "error")
            return None

    def _is_valid_docx(self, file_path: str) -> bool:
        """Verifica se é um arquivo .docx válido com estrutura interna"""
        if not file_path.lower().endswith('.docx'):
            return False
        try:
            with zipfile.ZipFile(file_path, 'r') as z:
                return 'word/document.xml' in z.namelist()
        except:
            return False

class Scheduler:
    """Fila de tarefas com execução sequencial"""

    def __init__(self):
        self.task_queue = queue.Queue()
        self.task_log = {}
        self.task_id_counter = 1
        self.executing = False
        self.lock = threading.Lock()

        # Start the background task processor
        self.processor_thread = threading.Thread(target=self._task_processor_loop, daemon=True)
        self.processor_thread.start()

    def schedule_task(self, task_func, task_args, run_time, description, dialog):
        """Schedules a task to run at a specific time"""
        with self.lock:
            task_id = self.task_id_counter
            self.task_id_counter += 1
            
            task_data = {
                "id": task_id,
                "func": task_func,
                "args": task_args,
                "run_time": run_time,
                "description": description,
                "status": "scheduled",
                "dialog": dialog  
            }
            
            self.task_log[task_id] = task_data
            self._enqueue_task(task_data)
            
            return task_id

               

    def _enqueue_task(self, task_data):
        task_data["status"] = "queued"
        self.task_queue.put(task_data)

    def _task_processor_loop(self):
        while True:
            task = self.task_queue.get()
            with self.lock:
                self.executing = True
                task["status"] = "waiting"
            
            now = datetime.now()
            if task["run_time"] > now:
                wait_seconds = (task["run_time"] - now).total_seconds()
                logging.info(f"Aguardando {wait_seconds} segundos para executar a tarefa {task['id']}")
                time.sleep(wait_seconds)

            with self.lock:
                task["status"] = "running"
            try:
                # Aqui está a correção principal - desempacota os argumentos corretamente
                if isinstance(task["args"], (tuple, list)):
                    task["func"](*task["args"])
                else:
                    task["func"](task["args"])
                    
                task["status"] = "completed"
            except Exception as e:
                task["status"] = f"error: {str(e)}"
                logging.exception(f"Erro ao executar tarefa agendada {task['id']}")
            finally:
                with self.lock:
                    self.executing = False

    def cancel_task(self, task_id: int):
        if task_id in self.task_log and self.task_log[task_id]["status"] == "scheduled":
            self.task_log[task_id]["status"] = "cancelled"

    def get_scheduled_tasks(self) -> Dict[int, Dict]:
        return self.task_log

class DocumentConverterApp:
    """Main application with optimized GUI"""

    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("DocWise V1 - Document Converter")

        self.template_manager = TemplateManager()

        # Configuração inicial
        self.config = EnhancedConfigManager()

        # ✅ Elimina processos presos do LibreOffice
        LibreOfficeDocument(
            libreoffice_path=self.config.config.get("libreoffice_path", "")
        )._kill_libreoffice_processes()

        self.scheduler = Scheduler()
        self.setup_ui()
        self.load_config()
        
        # Inicialização dos processadores
        self._initialize_processors()
        
        # Gerenciador de agendamento com múltiplos modelos
        self.task_manager = TaskManagerDeclaracoes(
            scheduler=self.scheduler,
            config=self.config,
            dialog=self.root,
        )

        self.current_processor = None
        self._setup_progress_throttle()
        self._initialize_notification_system()
        
        # Inicia os loops de processamento
        self._start_background_tasks()

        self.template_manager = TemplateManager()
        self.setup_template_library_menu()
        self.setup_batch_import_menu()

    def setup_batch_import_menu(self):
            menubar = self.root.nametowidget(".!menu")
            library_menu = menubar.winfo_children()[3]  # Ajuste o índice conforme seu menu
            # library_menu.add_command(
                # label="Importar Pasta...",
                # command=self.import_templates_batch
            # )

    def show_add_template_dialog(self):
        """Mostra diálogo para adicionar novo template"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Adicionar Template")
        dialog.geometry("600x300")

        # Campos do formulário
        ttk.Label(dialog, text="Nome do Template:").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        nome_entry = ttk.Entry(dialog, width=40)
        nome_entry.grid(row=0, column=1, columnspan=2, padx=5, pady=5)

        ttk.Label(dialog, text="Arquivo do Template:").grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        arquivo_entry = ttk.Entry(dialog, width=40)
        arquivo_entry.grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(dialog, text="Procurar", command=lambda: self.select_template_file(arquivo_entry)).grid(row=1, column=2)

        ttk.Label(dialog, text="Categoria:").grid(row=2, column=0, padx=5, pady=5, sticky=tk.W)
        categoria_entry = ttk.Entry(dialog, width=40)
        categoria_entry.grid(row=2, column=1, columnspan=2, padx=5, pady=5)

        ttk.Label(dialog, text="Tags (separadas por vírgula):").grid(row=3, column=0, padx=5, pady=5, sticky=tk.W)
        tags_entry = ttk.Entry(dialog, width=40)
        tags_entry.grid(row=3, column=1, columnspan=2, padx=5, pady=5)

        ttk.Label(dialog, text="Campos Obrigatórios (separados por vírgula):").grid(row=4, column=0, padx=5, pady=5, sticky=tk.W)
        campos_entry = ttk.Entry(dialog, width=40)
        campos_entry.grid(row=4, column=1, columnspan=2, padx=5, pady=5)

        ttk.Label(dialog, text="Descrição:").grid(row=5, column=0, padx=5, pady=5, sticky=tk.W)
        descricao_entry = ttk.Entry(dialog, width=40)
        descricao_entry.grid(row=5, column=1, columnspan=2, padx=5, pady=5)

        # Botões
        btn_frame = ttk.Frame(dialog)
        btn_frame.grid(row=6, column=0, columnspan=3, pady=10)
        
        ttk.Button(btn_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="Adicionar", command=lambda: self.add_template_to_library(
            nome_entry.get(),
            arquivo_entry.get(),
            categoria_entry.get(),
            [t.strip() for t in tags_entry.get().split(",")],
            [c.strip() for c in campos_entry.get().split(",")],
            descricao_entry.get(),
            dialog
        )).pack(side=tk.RIGHT)

    def select_template_file(self, entry_widget):
        """Abre diálogo para selecionar arquivo de template"""
        filepath = filedialog.askopenfilename(
            title="Selecionar Template",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        if filepath:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, filepath)

    def add_template_to_library(self, nome, arquivo, categoria, tags, campos_obrigatorios, descricao, dialog):
        """Adiciona template à biblioteca"""
        if not all([nome, arquivo, categoria]):
            messagebox.showerror("Erro", "Nome, arquivo e categoria são obrigatórios!")
            return
        
        try:
            # Verifica se o arquivo existe
            if not Path(arquivo).exists():
                messagebox.showerror("Erro", f"Arquivo não encontrado: {arquivo}")
                return

            self.template_manager.adicionar_template(
                nome=nome,
                arquivo=arquivo,
                categoria=categoria,
                tags=tags,
                campos_obrigatorios=campos_obrigatorios,
                descricao=descricao
            )
            messagebox.showinfo("Sucesso", "Template adicionado com sucesso!")
            dialog.destroy()
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao adicionar template:\n{str(e)}")
            logger.error(f"Erro ao adicionar template: {str(e)}")

    def import_templates_batch(self):
        folder = filedialog.askdirectory(title="Selecione a pasta com templates")
        if folder:
            importer = TemplateBatchImporter(self.template_manager)
            importer.import_from_folder(folder)
            messagebox.showinfo("Sucesso", "Templates importados em lote!")

    def setup_template_library_menu(self):
        menubar = self.root.nametowidget(".!menu")  # Acessa a barra de menus
        
        library_menu = tk.Menu(menubar, tearoff=0)
        library_menu.add_command(label="Abrir Biblioteca", 
                            command=self.show_template_library)
        library_menu.add_command(label="Adicionar Template", 
                            command=self.show_add_template_dialog)
        menubar.add_cascade(label="Biblioteca", menu=library_menu)

    def show_template_library(self):
        """Mostra a janela da biblioteca de templates"""
        dialog = TemplateLibraryDialog(
            parent=self.root,
            template_manager=self.template_manager,
            multi_select=False  # Altere para True se quiser seleção múltipla
        )
        self.root.wait_window(dialog)
        
        if dialog.selected_template:
            print(f"Template selecionado: {dialog.selected_template}")
          
            self.template_path.delete(0, tk.END)
            self.template_path.insert(0, str(dialog.selected_template))
        else:
            print("Nenhum template foi selecionado")

    def _initialize_processors(self):
        """Inicializa os processadores de documentos"""
        max_threads = self.config.config.get("max_threads", 4)
        use_libreoffice = self.config.config.get("use_libreoffice", False)
        libreoffice_path = self.config.config.get("libreoffice_path", "")
        ocr_enabled = self.config.config.get("ocr_enabled", False)
        tesseract_path = self.config.config.get("tesseract_path", "")
        
        self.declaration_gen = BatchDeclarationGenerator(
            max_workers=max_threads,
            use_libreoffice=use_libreoffice,
            libreoffice_path=libreoffice_path
        )
        
        self.docx_converter = DocumentConverter(
            max_workers=max_threads,
            use_libreoffice=use_libreoffice,
            libreoffice_path=libreoffice_path,
            ocr_enabled=ocr_enabled,
            tesseract_path=tesseract_path
        )
        
        self.pdf_converter = DocumentConverter(
            max_workers=max_threads,
            use_libreoffice=False,
            libreoffice_path="",
            ocr_enabled=ocr_enabled,
            tesseract_path=tesseract_path
        )

    def _setup_progress_throttle(self):
        """Configura o throttle de progresso baseado nas configurações"""
        throttle = self.config.config.get("progress_throttle", 0.1)
        self.declaration_gen.progress_throttle = throttle
        self.docx_converter.progress_throttle = throttle
        self.pdf_converter.progress_throttle = throttle

    def _initialize_notification_system(self):
        """Inicializa o sistema de notificações"""
        self.notification_window = None
        self.notification_queue = deque(maxlen=10)

    def _start_background_tasks(self):
        """Inicia os loops de processamento em background"""
        self.root.after(100, self.process_queues)
        self.root.after(1000, self.check_scheduled_tasks)

    # No seu DocumentConverterApp, onde agenda as tarefas:
    def schedule_declaration_task(self, excel_path: str, template_path: str, 
                                output_folder: str, create_zip: bool = False,
                                filename_pattern: str = None, 
                                save_in_subfolders: bool = False):
        """Usa o TaskManagerDeclaracoes para agendar"""
        if not filename_pattern:
            filename_pattern = self.filename_pattern_entry.get().strip()
        
        self.task_manager.agendar_declaracao(
            excel_path=excel_path,
            modelo_path=template_path,
            pasta_saida=output_folder,
            criar_zip=create_zip,
            filename_pattern=filename_pattern,  # Passa o padrão
            save_in_subfolders=save_in_subfolders
        )
    
    def setup_ui(self):
        style = ttk.Style()
        available_themes = style.theme_names()
        theme = self.config.config.get("theme", "clam")
        if theme in available_themes:
            style.theme_use(theme)
        
        self.root.geometry(self.config.config.get("window_size", "900x750"))
        
        self.setup_menu()
        
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.setup_declarations_tab()
        self.setup_conversion_tabs()
        self.setup_history_tab()
        
        self.setup_status_bar()
        
        self.notebook.select(self.config.config.get("last_tab", 0))
    
    def setup_menu(self):
        menubar = tk.Menu(self.root)
        
        # Arquivos menu
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Sair", command=self.on_closing)
        menubar.add_cascade(label="Arquivos", menu=file_menu)
        
        # bla vla menu
        config_menu = tk.Menu(menubar, tearoff=0)
        
        # temas submenu
        theme_menu = tk.Menu(config_menu, tearoff=0)
        style = ttk.Style()
        for theme in style.theme_names():
            theme_menu.add_command(
                label=theme,
                command=lambda t=theme: self.change_theme(t)
            )
        config_menu.add_cascade(label="Temas", menu=theme_menu)
        
        # Thread configuracao
        config_menu.add_command(
            label="Configurar Threads",
            command=self.configure_threads
        )
        
        # Notificacao config
        notification_menu = tk.Menu(config_menu, tearoff=0)
        self.enable_sounds_var = tk.BooleanVar(value=self.config.config.get("enable_sounds", True))
        notification_menu.add_checkbutton(
            label="Habilitar som de notificação",
            variable=self.enable_sounds_var,
            command=self.toggle_sound_notifications
        )
        self.show_notifications_var = tk.BooleanVar(value=self.config.config.get("show_notifications", True))
        notification_menu.add_checkbutton(
            label="Exibir notificações visuais",
            variable=self.show_notifications_var,
            command=self.toggle_visual_notifications
        )
        config_menu.add_cascade(label="Notificações", menu=notification_menu)
        
        # LibreOffice sei la 
        libreoffice_menu = tk.Menu(config_menu, tearoff=0)
        self.use_libreoffice_var = tk.BooleanVar(value=self.config.config.get("use_libreoffice", False))
        libreoffice_menu.add_checkbutton(
            label="Use LibreOffice para conversão",
            variable=self.use_libreoffice_var,
            command=self.toggle_libreoffice
        )
        libreoffice_menu.add_command(
            label="Configurar Caminho do LibreOffice",
            command=self.configure_libreoffice_path
        )
        config_menu.add_cascade(label="LibreOffice", menu=libreoffice_menu)
        
        # OCR 
        ocr_menu = tk.Menu(config_menu, tearoff=0)
        self.ocr_enabled_var = tk.BooleanVar(value=self.config.config.get("ocr_enabled", False))
        ocr_menu.add_checkbutton(
            label="ativar OCR",
            variable=self.ocr_enabled_var,
            command=self.toggle_ocr
        )
        ocr_menu.add_command(
            label="onfigurar o caminho do Tesseract OCR",
            command=self.configure_tesseract_path
        )
        config_menu.add_cascade(label="OCR Configurações", menu=ocr_menu)
        
        menubar.add_cascade(label="Configurações", menu=config_menu)
        
        # Tarefas menu 
        schedule_menu = tk.Menu(menubar, tearoff=0)
        schedule_menu.add_command(label="Agendar Tarefa Única", command=self.show_schedule_dialog)
        schedule_menu.add_command(label="Agendar Múltiplos Modelos", command=self.show_multi_model_dialog)
        schedule_menu.add_command(label="Ver Tarefas Agendadas", command=self.show_scheduled_tasks)
        menubar.add_cascade(label="Agendar", menu=schedule_menu)

        # Ajuda menu
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="Documentação", command=self.show_documentation)
        help_menu.add_command(label="Sobre", command=self.show_about)
        menubar.add_cascade(label="Ajuda", menu=help_menu)
        
        self.root.config(menu=menubar)
    
    def toggle_sound_notifications(self):
        self.config.config["Ativar sons"] = self.enable_sounds_var.get()
        self.config.save_config()
    
    def toggle_visual_notifications(self):
        self.config.config["Mostrar notificações"] = self.show_notifications_var.get()
        self.config.save_config()
    
    def toggle_ocr(self):
        self.config.config["Ativar OCR"] = self.ocr_enabled_var.get()
        self.config.save_config()
        self.docx_converter.ocr_enabled = self.ocr_enabled_var.get()
        self.pdf_converter.ocr_enabled = self.ocr_enabled_var.get()
        
        if self.ocr_enabled_var.get() and not self.config.config.get("tesseract_path", ""):
            self.configure_tesseract_path()
    
    def configure_tesseract_path(self):
        # Obtém o caminho atual do Tesseract salvo nas configurações, se existir
        current_path = self.config.config.get("tesseract_path", "")
        new_path = filedialog.askopenfilename(

            title="Selecione o executável do Tesseract OCR (tesseract.exe)",
            initialfile="tesseract.exe",
            filetypes=[("Executável do Tesseract", "tesseract.exe"), ("Todos os arquivos", "*.*")]
        )
        
        if new_path:
            self.config.config["tesseract_path"] = new_path
            self.config.save_config()
            self.docx_converter.tesseract_path = new_path
            self.pdf_converter.tesseract_path = new_path
            messagebox.showinfo("Successo", f"Caminho do Tesseract definido como:\n{new_path}")
    
    def toggle_libreoffice(self):
        current = not self.use_libreoffice_var.get()
        self.config.config["use_libreoffice"] = current
        self.config.save_config()
        self.docx_converter.use_libreoffice = current
        
        if current and not self.config.config.get("libreoffice_path", ""):
            self.configure_libreoffice_path()
    
    def configure_libreoffice_path(self):
        """Configura o caminho do executável do LibreOffice com verificação"""
        default_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice*\program\soffice.exe",  # Para versões com número
            r"C:\Program Files (x86)\LibreOffice*\program\soffice.exe"
        ]
        
        # Tenta encontrar automaticamente primeiro
        auto_path = ""
        for path in default_paths:
            matches = glob.glob(path)
            if matches:
                auto_path = matches[0]
                break
        
        # Diálogo para seleção manual
        new_path = filedialog.askopenfilename(
            title="Selecione o executável do LibreOffice (soffice.exe)",
            initialdir=os.path.dirname(auto_path) if auto_path else None,
            initialfile="soffice.exe",
            filetypes=[("Executável LibreOffice", "soffice.exe"), ("Todos os arquivos", "*.*")]
        )
        
        if new_path:
            # Verifica se o arquivo é válido
            if os.path.basename(new_path).lower() != "soffice.exe":
                messagebox.showerror("Erro", "Por favor, selecione o arquivo 'soffice.exe'")
                return
                
            # Testa a versão para verificar se é válido
            try:
                result = subprocess.run(
                    [new_path, "--version"],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=5,
                    text=True
                )
                
                if result.returncode != 0 or not result.stdout.strip():
                    raise RuntimeError("Falha ao verificar versão")
                    
                # Se chegou aqui, o caminho é válido
                self.config.config["libreoffice_path"] = new_path
                self.config.save_config()
                self.docx_converter.libreoffice_path = new_path
                messagebox.showinfo("Sucesso", f"Caminho do LibreOffice definido para:\n{new_path}")
                
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao validar o LibreOffice:\n{str(e)}\n"
                                        "Certifique-se que o LibreOffice está instalado corretamente.")
                logger.error(f"Erro ao validar LibreOffice: {str(e)}")
    
    def configure_threads(self):
        cpu_count = os.cpu_count() or 1
        current = self.config.config.get("max_threads", min(4, cpu_count))
        
        max_threads = simpledialog.askinteger(
            "Configure Threads",
            f"Maximum number of threads (1-{min(16, cpu_count * 2)})\n"
            f"Available CPUs: {cpu_count}\n"
            f"Recommended: {cpu_count} to {cpu_count * 2}",
            parent=self.root,
            minvalue=1,
            maxvalue=min(16, cpu_count * 2),
            initialvalue=current
        )
        
        if max_threads:
            self.config.config["max_threads"] = max_threads
            self.config.save_config()
            
            self.declaration_gen.max_workers = max_threads
            self.docx_converter.max_workers = max_threads
            self.pdf_converter.max_workers = max_threads
            
            messagebox.showinfo(
                "Configuração de Threads",
                f"Número máximo de threads definido como {max_threads}"
            )
    
    def change_theme(self, theme: str):
        style = ttk.Style()
        style.theme_use(theme)
        self.config.config["Temas"] = theme
        self.config.save_config()
    
    def show_schedule_dialog(self):
        """Exibe uma janela para agendar uma tarefa"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Agendar tarefa unica")
        dialog.geometry("400x300")
        dialog.resizable(False, False)
        
        ttk.Label(dialog, text="Agendar tarefa", font=('Arial', 12, 'bold')).pack(pady=10)
        
        # Task type selection
        ttk.Label(dialog, text="Tipo de tarefa:").pack(anchor=tk.W, padx=20)
        task_type = ttk.Combobox(dialog, values=["Gerar Declarações", "Converter DOCX para PDF", "Converter PDF para DOCX"])
        task_type.pack(fill=tk.X, padx=20, pady=5)
        task_type.current(0)
        
        # Schedule time
        ttk.Label(dialog, text="começa em:").pack(anchor=tk.W, padx=20)
        time_frame = ttk.Frame(dialog)
        time_frame.pack(fill=tk.X, padx=20, pady=5)
        
        hour_var = tk.StringVar(value=datetime.now().strftime("%H"))
        minute_var = tk.StringVar(value=datetime.now().strftime("%M"))
        
        ttk.Entry(time_frame, textvariable=hour_var, width=3).pack(side=tk.LEFT)
        ttk.Label(time_frame, text=":").pack(side=tk.LEFT)
        ttk.Entry(time_frame, textvariable=minute_var, width=3).pack(side=tk.LEFT)
        ttk.Label(time_frame, text="(24-houra formato)").pack(side=tk.LEFT, padx=5)
        
        # Date selection
        ttk.Label(dialog, text="Data:").pack(anchor=tk.W, padx=20)
        date_var = tk.StringVar(value=datetime.now().strftime("%Y-%m-%d"))
        ttk.Entry(dialog, textvariable=date_var).pack(fill=tk.X, padx=20, pady=5)
        
        # Description
        ttk.Label(dialog, text="Descrição:").pack(anchor=tk.W, padx=20)
        desc_entry = ttk.Entry(dialog)
        desc_entry.pack(fill=tk.X, padx=20, pady=5)
        
        # Buttons
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, pady=10, padx=20)
        
        ttk.Button(btn_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        ttk.Button(
        btn_frame, 
        text="Agendar", 
        command=lambda: self.schedule_task(
        task_type=task_type.get(),
        hour=hour_var.get(),
        minute=minute_var.get(),
        date=date_var.get(),
        description=desc_entry.get(),
        dialog=dialog
        )).pack(side=tk.RIGHT)

        task_type.get(),
        hour_var.get(),
        minute_var.get(),
        date_var.get(),
        desc_entry.get(),
        dialog
        
    def schedule_task(self, task_type: str, hour: str, minute: str, date: str, description: str, dialog: tk.Toplevel):
        """Agenda uma tarefa para ser executada em um horário específico"""
        try:
            # Validate time
            run_time = datetime.strptime(f"{date} {hour}:{minute}", "%Y-%m-%d %H:%M")
            if run_time < datetime.now():
                messagebox.showerror("Erro", "O horário agendado deve ser no futuro")
                return
            
            # Determine which function to schedule
            if task_type == "Gerar Declarações":
                if not self.validate_paths(
                    (self.excel_path, "Planilha Excel"),
                    (self.template_path, "Modelo do Word"),
                    (self.output_folder, "Pasta de saída")
                ):
                    return
                
                task_func = self._schedule_declaration_generation
                task_args = (
                    self.excel_path.get(),
                    self.template_path.get(),
                    self.output_folder.get(),
                    self.zip_output.get()
                )
            elif task_type == "Converter DOCX para PDF":
                if not self.validate_paths(
                    (self.docx_listbox, "Arquivos DOCX"),
                    (self.docx_output_folder, "Pasta de saída")
                ):
                    return
                
                task_func = self._schedule_docx_conversion
                task_args = (
                    list(self.docx_listbox.get(0, tk.END)),
                    self.docx_output_folder.get()
                )
            elif task_type == "Converter PDF para DOCX":
                if not self.validate_paths(
                    (self.pdf_listbox, "Arquivos PDF"),
                    (self.pdf_output_folder, "Pasta de saída")
                ):
                    return
                
                task_func = self._schedule_pdf_conversion
                task_args = (
                    list(self.pdf_listbox.get(0, tk.END)),
                    self.pdf_output_folder.get()
                )
            else:
                messagebox.showerror("Erro", "Tipo de tarefa inválido")
                return
            
            # Schedule the task
            task_id = self.scheduler.schedule_task(
                task_func=task_func,
                task_args=task_args,
                run_time=run_time,
                description=description or f"{task_type} at {run_time}",
                dialog=dialog  # Pass the dialog reference
            )
            
            messagebox.showinfo("Successo", f"Tarefa agendada com sucesso (ID: {task_id})")
            dialog.destroy()
        except ValueError:
            messagebox.showerror("Erro", "Formato de data ou hora inválido. Use o formato AAAA-MM-DD e 24 horas.")
                
    def _schedule_declaration_generation(self, excel_path: str, template_path: str, output_folder: str, create_zip: bool):
        "Wrapper para geração de declaração a ser chamado pelo planejador"
    
        processor = BatchDeclarationGenerator(self.config.config.get("max_threads", 4))
        processor.custom_filename_pattern = self.filename_pattern_entry.get().strip()
        processor.save_in_subfolders = self.save_in_subfolders_var.get()
        processor.subfolder_column = "Nome"
        
       
        self.current_processor = processor
        
       
        processor.generate_declarations(
            excel_path,
            template_path,
            self.output_folder.get(),
            create_zip,
            self.config
        )
        
    def _schedule_docx_conversion(self, docx_paths: List[str], output_folder: str):
        
        self.docx_listbox.delete(0, tk.END)
        for path in docx_paths:
            self.docx_listbox.insert(tk.END, path)
        self.docx_output_folder.delete(0, tk.END)
        self.docx_output_folder.insert(0, output_folder)
        
        self.notebook.select(1)  
        self.convert_docx_to_pdf()
    
    def _schedule_pdf_conversion(self, pdf_paths: List[str], output_folder: str):
       
        self.pdf_listbox.delete(0, tk.END)
        for path in pdf_paths:
            self.pdf_listbox.insert(tk.END, path)
        self.pdf_output_folder.delete(0, tk.END)
        self.pdf_output_folder.insert(0, output_folder)
        
        self.notebook.select(2)  
        self.convert_pdf_to_docx()
    
    def show_scheduled_tasks(self):
       
        tasks = self.scheduler.get_scheduled_tasks()
        
        dialog = tk.Toplevel(self.root)
        dialog.title("Tarefas agendadas")
        dialog.geometry("600x400")
        
        # Treeview to display tasks
        tree = ttk.Treeview(dialog, columns=("id", "descrição", "tempo de execução", "status"), show="títulos")
        tree.heading("id", text="ID")
        tree.heading("descrição", text="Descrição")
        tree.heading("tempo de execução", text="Horário agendado")
        tree.heading("status", text="Status")
        
        tree.column("id", width=50)
        tree.column("descrição", width=250)
        tree.column("tempo de execução", width=150)
        tree.column("status", width=100)
        
        for task_id, task_info in tasks.items():
            tree.insert("", "fim", values=(
                task_id,
                task_info["descrição"],
                task_info["tempo de execução"].strftime("%Y-%m-%d %H:%M"),
                task_info["status"]
            ))
        
        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Cancel button
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        ttk.Button(btn_frame, text="Cancelar Selecionado", command=lambda: self.cancel_selected_task(tree)).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Fechar", command=dialog.destroy).pack(side=tk.RIGHT)
    
    def cancel_selected_task(self, tree: ttk.Treeview):
        """Cancela a tarefa agendada selecionada"""
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("Aviso", "Nenhuma tarefa selecionada")
            return
        
        task_id = int(tree.item(selected[0], "valores")[0])
        self.scheduler.cancel_task(task_id)
        messagebox.showinfo("Successp", f"Tarefa {task_id} cancelada")
        
        
        self.show_scheduled_tasks()
    
    def check_scheduled_tasks(self):
        """Verifica periodicamente as tarefas agendadas (chamadas pelo loop principal)"""
        self.scheduler.get_scheduled_tasks()  
        self.root.after(60000, self.check_scheduled_tasks)  
    
    def show_documentation(self):
        doc_window = tk.Toplevel(self.root)
        doc_window.title("DocWise Documentos")
        doc_window.geometry("800x600")
        
        notebook = ttk.Notebook(doc_window)
        notebook.pack(fill=tk.BOTH, expand=True)
        
        # Declaration generation guide
        decl_frame = ttk.Frame(notebook)
        notebook.add(decl_frame, text="Geração de Declaração")
        
        decl_text = ScrolledText(decl_frame, wrap=tk.WORD, font=('Arial', 10))
        decl_text.pack(fill=tk.BOTH, expand=True)
        
        decl_guide = """
        ═══════════════════════════════
          📰 GERADOR DE DECLARAÇÕES
        ═══════════════════════════════

        ▼▼▼ PASSO A PASSO ▼▼▼

        ▣ 1. PLANILHA EXCEL:
        → Selecione o arquivo com os dados
        → Verifique se os títulos das colunas batem com:
            {{NOME}}, {{MATRICULA}}, {{DATA}} etc.
        ✚ Opcional: Use "Detectar Campos" para validação

        ▣ 2. MODELO WORD:
        → Use {{CAMPOS}} nos locais de preenchimento
        → Visualize com "Pré-visualizar"
        ✚ Dica: Salve como .docx (não .doc)

        ▣ 3. CONFIGURAÇÕES:
        → Pasta de saída (evite caminhos longos)
        → Padrão de nomes: (Tipo)_Nome.docx
        ✓ Opção ZIP para múltiplos arquivos

        ▣ 4. GERAR:
        → Barra de progresso em tempo real
        → Log detalhado de erros
        → Relatório automático em PDF/CSV

        ⚠️ ERROS COMUNS:
        × Campos não preenchidos:
            → Verifique espaços em {{ NOME }} vs {{NOME}}
            → Confira maiúsculas/minúsculas

        × Problemas com formatação:
            → Evite tabelas/templates muito complexos
            → Teste com modelo simplificado primeiro

        ⚡ TURBO-DICAS:
        • Para 100+ arquivos: Agende para horário noturno
        • Nomes duplicados: São automaticamente numerados
        • Log completo em: /Logs/geracao_[data].txt
        """
        decl_text.insert(tk.END, decl_guide)
        decl_text.config(state=tk.DISABLED)
        
        # Conversion guide
        conv_frame = ttk.Frame(notebook)
        notebook.add(conv_frame, text="Conversão de documentos")
        
        conv_text = ScrolledText(conv_frame, wrap=tk.WORD, font=('Arial', 10))
        conv_text.pack(fill=tk.BOTH, expand=True)
        
        conv_guide = """
        ═══════════════════════════════
          📑 CONVERSÃO DE DOCUMENTOS
        ═══════════════════════════════

        ▸ DOCX para PDF:
        1. Selecione arquivos .docx
        2. Escolha pasta de saída
        3. Clique em "Converter"
            → LibreOffice: Melhor para tabelas/complexos

        ▸ PDF para DOCX:
        1. Selecione arquivos .pdf
        2. Escolha pasta de saída
        3. Clique em "Converter"
            → OCR: Ative para PDFs escaneados

        ⚠️ PROBLEMAS COMUNS (E SOLUÇÕES):
        • Erro de caracteres especiais (ç, á, ñ):
            → Use nomes sem acentos nos arquivos
            → Prefira o Microsoft Word se possível

        • Arquivo não convertido:
            → Verifique se não está aberto em outro programa
            → Tente converter individualmente

        • LibreOffice travando:
            → Reinicie o programa
            → Atualize sua versão do LibreOffice

        • Arquivos duplicados:
            → São renomeados automaticamente (arquivo_1.pdf)

        ⚡ DICAS PRO:
        • Para muitos arquivos: Agende conversões noturnas
        • Problemas persistentes? Exporte o log e envie ao suporte
        • Relatórios completos são gerados em /Relatorios/
        """
        conv_text.insert(tk.END, conv_guide)
        conv_text.config(state=tk.DISABLED)
        
        # History guide
        hist_frame = ttk.Frame(notebook)
        notebook.add(hist_frame, text="Histórico de Execução")
        
        hist_text = ScrolledText(hist_frame, wrap=tk.WORD, font=('Arial', 10))
        hist_text.pack(fill=tk.BOTH, expand=True)
        
        hist_guide = """
        ═══════════════════════════════
          📜 HISTÓRICO DE EXECUÇÕES
        ═══════════════════════════════

        📋 O QUE É ARMAZENADO:
        • Geração de declarações
        • Conversões de documentos
        • Tarefas agendadas
        • Processos manuais

        📊 DETALHES REGISTRADOS:
        ├─ Data/hora exata da operação
        ├─ Tipo: Conversão/Declaração/Agendamento
        ├─ Arquivos de entrada utilizados
        ├─ Pasta de saída dos resultados
        ├─ Status: Completo/Falha/Interrompido
        ├─ Itens processados vs erros
        └─ Tempo total de execução

        🔍 COMO UTILIZAR:
        ▸ Filtrar por:
            - Tipo de operação (dropdown)
            - Período (calendário)
            - Status (completo/falha)

        ▸ Visualizar detalhes:
            - Duplo-clique em qualquer registro
            - Painel expandido com:
                * Caminhos completos
                * Erros específicos
                * Log técnico

        💾 EXPORTAR DADOS:
        ✓ CSV para análise externa
        ✓ PDF com relatório consolidado
        ✓ Copiar dados para clipboard

        ⚠️ ATENÇÃO:
        • Histórico limitado aos últimos 500 registros
        • Dados antigos são automaticamente compactados
        • Use exportação para backup importante
        """
        hist_text.insert(tk.END, hist_guide)
        hist_text.config(state=tk.DISABLED)
        
        btn_frame = ttk.Frame(doc_window)
        btn_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frame, text="Export Log", command=self.export_log).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Fechar", command=doc_window.destroy).pack(side=tk.RIGHT)
    
    def export_log(self):
        
        log_file = filedialog.asksaveasfilename(
            title="Salvar log de execução",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("All files", "*.*")]
        )
        
        if not log_file:
            return
        
        try:
            with open(log_file, 'w', encoding='utf-8') as f:
                # Write system log
                with open('document_converter.log', 'r', encoding='utf-8') as log_f:
                    f.write("=== SYSTEM LOG ===\n")
                    f.write(log_f.read())
                    f.write("\n\n")
                
                # Write interface log
                current_log = self.get_current_log_widget()
                if current_log:
                    f.write("=== INTERFACE LOG ===\n")
                    f.write(current_log.get("1.0", tk.END))
            
            messagebox.showinfo("Success", f"Log exported to:\n{log_file}")
            # Open file in default viewer
            webbrowser.open(log_file)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export log: {str(e)}")
    
    def show_about(self):
        about_window = tk.Toplevel(self.root)
        about_window.title("Sobre o DocWise")
        about_window.geometry("500x400")
        about_window.resizable(False, False)

        main_frame = ttk.Frame(about_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        ttk.Label(
            main_frame,
            text="DocWise - Document Converter",
            font=('Arial', 14, 'bold')
        ).pack(pady=10)

        info_frame = ttk.Frame(main_frame)
        info_frame.pack(fill=tk.X, pady=5)

        # Informações do sistema
        ttk.Label(info_frame, text="Versão:", width=10, anchor=tk.W).pack(side=tk.LEFT)
        ttk.Label(info_frame, text="2.0", anchor=tk.W).pack(side=tk.LEFT, fill=tk.X, expand=True)

        ttk.Label(info_frame, text="Autor:", width=10, anchor=tk.W).pack(side=tk.LEFT)
        ttk.Label(info_frame, text="Erlon Lopes", anchor=tk.W).pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Mensagem inspiradora
        ttk.Label(
            main_frame,
            text="\nMensagem:",
            font=('Arial', 10, 'bold')
        ).pack(anchor=tk.W)

        ttk.Label(
            main_frame,
            text='"Não temas, porque eu sou contigo; '
                'não te assombres, porque eu sou o teu Deus."\n- Isaías 41:10',
            wraplength=460,
            font=('Arial', 10, 'italic'),
            foreground="#333366",
            anchor=tk.CENTER,
            justify=tk.CENTER
        ).pack(pady=10)

        # Informações adicionais
        additional_info = [
            "Projeto open-source criado para estudo e automatização de tarefas.",
            "Desenvolvido com apoio de tecnologias modernas, IA, pesquisas, vídeos do YouTube e projetos anteriores.",
            "Código livre para uso, modificação e distribuição.",
            "Dúvidas ou sugestões? Entre em contato:",
            "erlon.araujo@upe.br | erlon.lopes2006@gmail.com",
        ]


        for line in additional_info:
            ttk.Label(main_frame, text=line, anchor=tk.W, wraplength=460, justify=tk.LEFT).pack(fill=tk.X)

        ttk.Button(main_frame, text="Fechar", command=about_window.destroy).pack(pady=20)
   
    def setup_declarations_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Gerar Declarações")
        
        # 🔹 Área dos arquivos
        file_frame = ttk.LabelFrame(tab, text="Arquivos", padding="10")
        file_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(file_frame, text="Planilha do Excel:").grid(row=0, column=0, sticky=tk.W)
        self.excel_path = ttk.Entry(file_frame, width=50)
        self.excel_path.grid(row=0, column=1, padx=5)
        ttk.Button(file_frame, text="Procurar", command=self.select_excel).grid(row=0, column=2)

        ttk.Label(file_frame, text="Modelo do Word:").grid(row=1, column=0, sticky=tk.W)
        self.template_path = ttk.Entry(file_frame, width=50)
        self.template_path.grid(row=1, column=1, padx=5)
        ttk.Button(file_frame, text="Procurar", command=self.select_template).grid(row=1, column=2)

        ttk.Label(file_frame, text="Pasta de saída:").grid(row=2, column=0, sticky=tk.W)
        self.output_folder = ttk.Entry(file_frame, width=50)
        self.output_folder.grid(row=2, column=1, padx=5)
        ttk.Button(file_frame, text="Procurar", command=self.select_output_folder).grid(row=2, column=2)
        
        # 🔹 Área das opções de geração de declarações
        declarations_frame = ttk.LabelFrame(tab, text="Opções de Declaração", padding="10")
        declarations_frame.pack(fill=tk.X, pady=5)

        # Padrão de nome dos arquivos
        self.filename_pattern_label = ttk.Label(declarations_frame, text="Modelo de nome dos arquivos:")
        self.filename_pattern_label.grid(row=2, column=0, sticky="w", padx=5, pady=5)

        self.filename_pattern_entry = ttk.Entry(declarations_frame)
        self.filename_pattern_entry.grid(row=2, column=1, sticky="ew", padx=5, pady=5)
        self.filename_pattern_entry.insert(0, "{{Nome}}_{{TipoDoc}}")  # Sugestão padrão mais clara

        # Checkbox salvar em subpastas
        self.save_in_subfolders_var = tk.BooleanVar()
        self.save_in_subfolders_check = ttk.Checkbutton(
            declarations_frame,
            text="Salvar arquivos em subpastas (por aluno)",
            variable=self.save_in_subfolders_var
        )
        self.save_in_subfolders_check.grid(row=3, column=0, columnspan=2, sticky="w", padx=5, pady=5)
        
        # 🔹 Opções adicionais (ZIP, Detectar campos, Preview)
        options_frame = ttk.Frame(file_frame)
        options_frame.grid(row=3, column=1, columnspan=2, sticky=tk.W, pady=5)

        self.zip_output = tk.BooleanVar(value=self.config.config.get("saída zip", False))
        ttk.Checkbutton(options_frame, text="Criar arquivo ZIP", variable=self.zip_output).pack(side=tk.LEFT)

        ttk.Button(options_frame, text="Detectar campos", command=self.detect_template_fields).pack(side=tk.LEFT, padx=10)
        ttk.Button(options_frame, text="Visualização", command=self.show_preview).pack(side=tk.LEFT)
        
        # 🔹 Área de log
        log_frame = ttk.LabelFrame(tab, text="Execution Log", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        self.declaration_log = ScrolledText(log_frame, height=15, wrap=tk.WORD)
        self.declaration_log.pack(fill=tk.BOTH, expand=True)

        self.declaration_log.tag_config("error", foreground="red")
        self.declaration_log.tag_config("warning", foreground="orange")
        self.declaration_log.tag_config("success", foreground="green")
        self.declaration_log.tag_config("info", foreground="blue")

        # 🔹 Botões de ação
        button_frame = ttk.Frame(tab)
        button_frame.pack(fill=tk.X, pady=5)

        ttk.Button(button_frame, text="Export Log", command=self.export_log).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="Gerar Declarações", command=self.start_declaration_generation).pack(side=tk.RIGHT)
        ttk.Button(button_frame, text="Parar", command=self.stop_process).pack(side=tk.RIGHT, padx=5)

    
    def setup_conversion_tabs(self):
        docx_tab = ttk.Frame(self.notebook)
        self.notebook.add(docx_tab, text="DOCX para PDF")
        self.setup_conversion_tab(docx_tab, "DOCX", "PDF", "docx_path", "docx_output_folder")
        
        pdf_tab = ttk.Frame(self.notebook)
        self.notebook.add(pdf_tab, text="PDF para DOCX")
        self.setup_conversion_tab(pdf_tab, "PDF", "DOCX", "pdf_path", "pdf_output_folder")
    
    def setup_conversion_tab(self, tab: ttk.Frame, from_ext: str, to_ext: str, 
                           path_attr: str, output_attr: str):
        file_frame = ttk.LabelFrame(tab, text="Arquivos", padding="10")
        file_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(file_frame, text=f"{from_ext} Arquivo(s):").grid(row=0, column=0, sticky=tk.W)
        
        list_frame = ttk.Frame(file_frame)
        list_frame.grid(row=1, column=0, columnspan=3, sticky=tk.EW, pady=5)
        
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        listbox = tk.Listbox(
            list_frame, 
            yscrollcommand=scrollbar.set,
            selectmode=tk.EXTENDED,
            height=4,
            width=60
        )
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=listbox.yview)
        
        setattr(self, f"{from_ext.lower()}_listbox", listbox)
        
        btn_frame = ttk.Frame(file_frame)
        btn_frame.grid(row=2, column=0, columnspan=3, sticky=tk.W)
        
        ttk.Button(btn_frame, text="Adicionar arquivo(s)", 
                  command=getattr(self, f"select_{path_attr}")).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Adicionar pasta", 
                  command=getattr(self, f"select_{path_attr}_folder")).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Remover Selecionado", 
                  command=getattr(self, f"remove_selected_{from_ext.lower()}")).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Limpar lista", 
                  command=getattr(self, f"clear_{from_ext.lower()}_list")).pack(side=tk.LEFT, padx=5)
        
        ttk.Label(file_frame, text="Pasta de saída:").grid(row=3, column=0, sticky=tk.W)
        output_entry = ttk.Entry(file_frame, width=50)
        output_entry.grid(row=3, column=1, padx=5)
        setattr(self, output_attr, output_entry)
        
        ttk.Button(file_frame, text="Navegar", 
                  command=getattr(self, f"select_{output_attr}")).grid(row=3, column=2)
        
        log_frame = ttk.LabelFrame(tab, text="Log de Execução", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        log_widget = ScrolledText(log_frame, height=15, wrap=tk.WORD)
        log_widget.pack(fill=tk.BOTH, expand=True)
        setattr(self, f"{from_ext.lower()}_log", log_widget)
        
        log_widget.tag_config("error", foreground="red")
        log_widget.tag_config("warning", foreground="orange")
        log_widget.tag_config("success", foreground="green")
        log_widget.tag_config("info", foreground="blue")
        
        button_frame = ttk.Frame(tab)
        button_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(button_frame, text="Export Log", command=self.export_log).pack(side=tk.LEFT)
        ttk.Button(button_frame, text=f"Converter para {to_ext}", 
                  command=getattr(self, f"convert_{from_ext.lower()}_to_{to_ext.lower()}")).pack(side=tk.RIGHT)
        ttk.Button(button_frame, text="Parar", 
                  command=self.stop_process).pack(side=tk.RIGHT, padx=5)
    
    def setup_history_tab(self):
         tab = ttk.Frame(self.notebook)
         self.notebook.add(tab, text="")
        

    #     filter_frame = ttk.LabelFrame(tab, text="Filters", padding="10")
    #     filter_frame.pack(fill=tk.X, pady=5)
        
    #     ttk.Label(filter_frame, text="Operation Type:").grid(row=0, column=0, sticky=tk.W)
    #     self.history_filter_type = ttk.Combobox(filter_frame, values=["All", "Declaration Generation", "DOCX to PDF", "PDF to DOCX"])
    #     self.history_filter_type.grid(row=0, column=1, sticky=tk.W, padx=5)
    #     self.history_filter_type.current(0)
        
    #     ttk.Label(filter_frame, text="Date Range:").grid(row=0, column=2, sticky=tk.W, padx=10)
    #     self.history_filter_from = ttk.Entry(filter_frame, width=10)
    #     self.history_filter_from.grid(row=0, column=3, sticky=tk.W)
    #     ttk.Label(filter_frame, text="to").grid(row=0, column=4, sticky=tk.W, padx=2)
    #     self.history_filter_to = ttk.Entry(filter_frame, width=10)
    #     self.history_filter_to.grid(row=0, column=5, sticky=tk.W)
        
    #     ttk.Button(filter_frame, text="Apply Filters", command=self.refresh_history).grid(row=0, column=6, padx=10)
    #     ttk.Button(filter_frame, text="Export History", command=self.export_history).grid(row=0, column=7)
        

    #     list_frame = ttk.Frame(tab)
    #     list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
    #     columns = ("id", "timestamp", "operation", "processed", "errors", "duration", "status")
    #     self.history_tree = ttk.Treeview(
    #         list_frame, 
    #         columns=columns,
    #         show="headings",
    #         selectmode="browse"
    #     )
        
    #     self.history_tree.heading("id", text="ID")
    #     self.history_tree.heading("timestamp", text="Timestamp")
    #     self.history_tree.heading("operation", text="Operation")
    #     self.history_tree.heading("processed", text="Processed")
    #     self.history_tree.heading("errors", text="Errors")
    #     self.history_tree.heading("duration", text="Duration")
    #     self.history_tree.heading("status", text="Status")
        
    #     self.history_tree.column("id", width=50, anchor=tk.CENTER)
    #     self.history_tree.column("timestamp", width=150)
    #     self.history_tree.column("operation", width=150)
    #     self.history_tree.column("processed", width=80, anchor=tk.CENTER)
    #     self.history_tree.column("errors", width=80, anchor=tk.CENTER)
    #     self.history_tree.column("duration", width=100, anchor=tk.CENTER)
    #     self.history_tree.column("status", width=100, anchor=tk.CENTER)
        
    #     scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.history_tree.yview)
    #     self.history_tree.configure(yscrollcommand=scrollbar.set)
        
    #     self.history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    #     scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        

    #     detail_frame = ttk.LabelFrame(tab, text="Details", padding="10")
    #     detail_frame.pack(fill=tk.BOTH, pady=5)
        
    #     self.history_details = ScrolledText(detail_frame, height=8, wrap=tk.WORD)
    #     self.history_details.pack(fill=tk.BOTH, expand=True)
        

    #     self.history_tree.bind("<<TreeviewSelect>>", self.show_history_details)
        

    #     self.refresh_history()
    
    def refresh_history(self):
  
      
        for item in self.history_tree.get_children():
            self.history_tree.delete(item)
        
       
        op_filter = self.history_filter_type.get()
        date_from = self.history_filter_from.get()
        date_to = self.history_filter_to.get()
        
       
        records = self.config.history.get_recent_history(self.config.config.get("max_history_items", 50))
        
      
        filtered_records = []
        for record in records:
           
            if op_filter != "All" and record["operation_type"] != op_filter:
                continue
            
          
            record_date = datetime.strptime(record["timestamp"], "%Y-%m-%d %H:%M:%S").date()
            
            if date_from:
                try:
                    from_date = datetime.strptime(date_from, "%Y-%m-%d").date()
                    if record_date < from_date:
                        continue
                except ValueError:
                    pass
            
            if date_to:
                try:
                    to_date = datetime.strptime(date_to, "%Y-%m-%d").date()
                    if record_date > to_date:
                        continue
                except ValueError:
                    pass
            
            filtered_records.append(record)
        
        
        for record in filtered_records:
            timestamp = datetime.strptime(record["timestamp"], "%Y-%m-%d %H:%M:%S")
            duration = f"{float(record['duration']):.2f}s" if record['duration'] else "N/A"
            
            self.history_tree.insert("", "end", values=(
                record["id"],
                timestamp.strftime("%Y-%m-%d %H:%M"),
                record["operation_type"],
                record["processed_count"],
                record["error_count"],
                duration,
                record["status"]
            ))
    
    def show_history_details(self, event):
        """Shows details for the selected history item"""
        selected = self.history_tree.selection()
        if not selected:
            return
        
        item = self.history_tree.item(selected[0])
        record_id = item["values"][0]
        
     
        with sqlite3.connect(self.config.history.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM execution_history WHERE id = ?", (record_id,))
            record = cursor.fetchone()
        
        if not record:
            self.history_details.config(state=tk.NORMAL)
            self.history_details.delete(1.0, tk.END)
            self.history_details.insert(tk.END, "Details not available")
            self.history_details.config(state=tk.DISABLED)
            return
        
  
        details = f"Operation: {record['operation_type']}\n"
        details += f"Timestamp: {record['timestamp']}\n"
        details += f"Status: {record['status']}\n"
        details += f"Processed: {record['processed_count']}\n"
        details += f"Errors: {record['error_count']}\n"
        
        if record['duration']:
            details += f"Duration: {float(record['duration']):.2f} seconds\n"
        
        details += f"\nInput Files:\n{record['input_files']}\n"
        details += f"\nOutput Folder:\n{record['output_folder']}\n"
        details += f"\nLog Summary:\n{record['log_text'][:1000]}..."
        
        self.history_details.config(state=tk.NORMAL)
        self.history_details.delete(1.0, tk.END)
        self.history_details.insert(tk.END, details)
        self.history_details.config(state=tk.DISABLED)
    
    def export_history(self):
        """Exports history to CSV file"""
        csv_file = filedialog.asksaveasfilename(
            title="Save History as CSV",
            defaultextension=".csv",
            filetypes=[("CSV Files", "*.csv"), ("All files", "*.*")]
        )
        
        if not csv_file:
            return
        
        try:
           
            op_filter = self.history_filter_type.get()
            date_from = self.history_filter_from.get()
            date_to = self.history_filter_to.get()
            
          
            records = self.config.history.get_recent_history(self.config.config.get("max_history_items", 50))
       
            filtered_records = []
            for record in records:
             
                if op_filter != "All" and record["operation_type"] != op_filter:
                    continue
                
             
                record_date = datetime.strptime(record["timestamp"], "%Y-%m-%d %H:%M:%S").date()
                
                if date_from:
                    try:
                        from_date = datetime.strptime(date_from, "%Y-%m-%d").date()
                        if record_date < from_date:
                            continue
                    except ValueError:
                        pass
                
                if date_to:
                    try:
                        to_date = datetime.strptime(date_to, "%Y-%m-%d").date()
                        if record_date > to_date:
                            continue
                    except ValueError:
                        pass
                
                filtered_records.append(record)
            
        
            with open(csv_file, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
             
                writer.writerow([
                    "ID", "Timestamp", "Operation Type", "Input Files", 
                    "Output Folder", "Status", "Processed Count", 
                    "Error Count", "Duration", "Log Summary"
                ])
                
            
                for record in filtered_records:
                    writer.writerow([
                        record["id"],
                        record["timestamp"],
                        record["operation_type"],
                        record["input_files"],
                        record["output_folder"],
                        record["status"],
                        record["processed_count"],
                        record["error_count"],
                        record["duration"],
                        record["log_text"][:1000]  # Truncate long logs
                    ])
            
            messagebox.showinfo("Success", f"History exported to:\n{csv_file}")
          
            webbrowser.open(csv_file)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export history: {str(e)}")
    

    def setup_status_bar(self):
        self.status_frame = ttk.Frame(self.root)
        self.status_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.status_label = ttk.Label(self.status_frame, text="Preparar")
        self.status_label.pack(side=tk.LEFT)
        
        progress_frame = ttk.Frame(self.status_frame)
        progress_frame.pack(side=tk.RIGHT, fill=tk.X, expand=True)
        
        self.progress = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.progress_label = ttk.Label(progress_frame, text="0%", width=5)
        self.progress_label.pack(side=tk.LEFT, padx=5)
        
        # Indeterminate progress bar for quick operations
        self.indeterminate_progress = ttk.Progressbar(
            progress_frame, 
            mode='indeterminate',
            length=100
        )
    
    def show_indeterminate_progress(self, show: bool = True):
        
        if show:
            self.progress.pack_forget()
            self.progress_label.pack_forget()
            self.indeterminate_progress.pack(side=tk.LEFT, fill=tk.X, expand=True)
            self.indeterminate_progress.start(10)
        else:
            self.indeterminate_progress.stop()
            self.indeterminate_progress.pack_forget()
            self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True)
            self.progress_label.pack(side=tk.LEFT, padx=5)
    
    def load_config(self):
        recent = self.config.config["recent_files"]
        self.excel_path.insert(0, recent.get("excel", ""))
        self.template_path.insert(0, recent.get("template", ""))
        self.output_folder.insert(0, recent.get("output", ""))
        
        if self.config.config.get("auto_load_last_files", True):
            for path in recent.get("docx", []):
                if os.path.exists(path):
                    self.docx_listbox.insert(tk.END, path)
            
            self.docx_output_folder.insert(0, recent.get("docx_output", ""))
            
            for path in recent.get("pdf", []):
                if os.path.exists(path):
                    self.pdf_listbox.insert(tk.END, path)
            
            self.pdf_output_folder.insert(0, recent.get("pdf_output", ""))
        
        self.zip_output.set(self.config.config.get("zip_output", False))
    
    def save_config(self):
        self.config.config["recent_files"] = {
            "excel": self.excel_path.get(),
            "template": self.template_path.get(),
            "output": self.output_folder.get(),
            "docx": list(self.docx_listbox.get(0, tk.END)),
            "docx_output": self.docx_output_folder.get(),
            "pdf": list(self.pdf_listbox.get(0, tk.END)),
            "pdf_output": self.pdf_output_folder.get()
        }
        
        self.config.config["zip_output"] = self.zip_output.get()
        self.config.config["window_size"] = self.root.geometry()
        self.config.config["last_tab"] = self.notebook.index(self.notebook.select())
        self.config.save_config()
    
    def select_excel(self):
        file_path = filedialog.askopenfilename(
            title="Selecione Planilha do Excel",
            filetypes=[("Arquivos do Excel", "*.xlsx;*.xls"), ("Todos os arquivos", "*.*")]
        )
        if file_path:
            self.excel_path.delete(0, tk.END)
            self.excel_path.insert(0, file_path)
            self.config.update_recent_file("excel", file_path)
    
    def select_template(self):
        file_path = filedialog.askopenfilename(
            title="Selecione o modelo do Word",
            filetypes=[("Arquivos do Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        if file_path:
            self.template_path.delete(0, tk.END)
            self.template_path.insert(0, file_path)
            self.config.update_recent_file("modelos", file_path)
    
    def select_output_folder(self):
        folder_path = filedialog.askdirectory(title="Selecione a pasta de saída")
        if folder_path:
            self.output_folder.delete(0, tk.END)
            self.output_folder.insert(0, folder_path)
            self.config.update_recent_file("saída", folder_path)
    
    def select_docx_path(self):
        file_paths = filedialog.askopenfilenames(
            title="Selecione arquivo(s) DOCX",
            filetypes=[("Arquivos do Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )
        if file_paths:
            for path in file_paths:
                if path not in self.docx_listbox.get(0, tk.END):
                    self.docx_listbox.insert(tk.END, path)
            self.config.update_recent_file("docx", list(self.docx_listbox.get(0, tk.END)))
    
    def select_docx_path_folder(self):
        folder_path = filedialog.askdirectory(title="Selecione a pasta com os arquivos DOCX")
        if folder_path:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    if file.lower().endswith('.docx'):
                        path = os.path.join(root, file)
                        if path not in self.docx_listbox.get(0, tk.END):
                            self.docx_listbox.insert(tk.END, path)
            self.config.update_recent_file("docx", list(self.docx_listbox.get(0, tk.END)))
    
    def remove_selected_docx(self):
        selected = self.docx_listbox.curselection()
        for i in reversed(selected):
            self.docx_listbox.delete(i)
        self.config.update_recent_file("docx", list(self.docx_listbox.get(0, tk.END)))
    
    def clear_docx_list(self):
        self.docx_listbox.delete(0, tk.END)
        self.config.update_recent_file("docx", [])
    
    def select_docx_output_folder(self):
        folder_path = filedialog.askdirectory(title="Selecione a pasta de saída para PDF")
        if folder_path:
            self.docx_output_folder.delete(0, tk.END)
            self.docx_output_folder.insert(0, folder_path)
            self.config.update_recent_file("docx_output", folder_path)
    
    def select_pdf_path(self):
        file_paths = filedialog.askopenfilenames(
            title="Selecione o(s) arquivo(s) PDF",
            filetypes=[("Arquivos PDF", "*.pdf"), ("Todos os arquivos", "*.*")]
        )
        if file_paths:
            for path in file_paths:
                if path not in self.pdf_listbox.get(0, tk.END):
                    self.pdf_listbox.insert(tk.END, path)
            self.config.update_recent_file("pdf", list(self.pdf_listbox.get(0, tk.END)))
    
    def select_pdf_path_folder(self):
        folder_path = filedialog.askdirectory(title="Selecione a pasta com os arquivos PDF")
        if folder_path:
            for root, _, files in os.walk(folder_path):
                for file in files:
                    if file.lower().endswith('.pdf'):
                        path = os.path.join(root, file)
                        if path not in self.pdf_listbox.get(0, tk.END):
                            self.pdf_listbox.insert(tk.END, path)
            self.config.update_recent_file("pdf", list(self.pdf_listbox.get(0, tk.END)))
    
    def remove_selected_pdf(self):
        selected = self.pdf_listbox.curselection()
        for i in reversed(selected):
            self.pdf_listbox.delete(i)
        self.config.update_recent_file("pdf", list(self.pdf_listbox.get(0, tk.END)))
    
    def clear_pdf_list(self):
        self.pdf_listbox.delete(0, tk.END)
        self.config.update_recent_file("pdf", [])
    
    def select_pdf_output_folder(self):
        folder_path = filedialog.askdirectory(title="Selecione a pasta de saída para DOCX")
        if folder_path:
            self.pdf_output_folder.delete(0, tk.END)
            self.pdf_output_folder.insert(0, folder_path)
            self.config.update_recent_file("pdf_output", folder_path)
    
    def detect_template_fields(self):
        template_path = self.template_path.get()
        if not template_path:
            messagebox.showwarning("Aviso", "Nenhum modelo selecionado")
            return
        
        self.show_indeterminate_progress(True)
        self.update_status("Analyzing template...")
        
        try:
            placeholders = self.declaration_gen.detect_placeholders(template_path)
            if not placeholders:
                messagebox.showinfo("Informações", "Nenhum campo detectado no modelo (use {{FIELD}})")
                return
            
            excel_path = self.excel_path.get()
            if excel_path and os.path.exists(excel_path):
                valid, validation_result = self.declaration_gen.validate_columns(excel_path, template_path, self.config)
            else:
                valid, validation_result = False, {
                    "mapping": {},
                    "missing": [],
                    "warnings": [],
                    "suggestions": {},
                    "columns": [],
                    "saved_mapping_used": False
                }
            
            fields_window = tk.Toplevel(self.root)
            fields_window.title("Relatório de Campo")
            fields_window.geometry("600x500")
            
            notebook = ttk.Notebook(fields_window)
            notebook.pack(fill=tk.BOTH, expand=True)
            
            fields_tab = ttk.Frame(notebook)
            notebook.add(fields_tab, text="Campos detectados")
            
            ttk.Label(fields_tab, text="Campos obrigatórios no modelo:", 
                     font=('Arial', 10, 'bold')).pack(pady=5, anchor=tk.W)
            
            fields_text = ScrolledText(fields_tab, height=10, wrap=tk.WORD)
            fields_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
            
            for field in placeholders:
                clean_field = field[2:-2]
                fields_text.insert(tk.END, f"• {clean_field}\n")
            
            fields_text.config(state=tk.DISABLED)
            
            if excel_path and os.path.exists(excel_path):
                validation_tab = ttk.Frame(notebook)
                notebook.add(validation_tab, text="Excel Validation")
                
                ttk.Label(validation_tab, text="Validation results:", 
                         font=('Arial', 10, 'bold')).pack(pady=5, anchor=tk.W)
                
                validation_text = ScrolledText(validation_tab, height=15, wrap=tk.WORD)
                validation_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
                
                validation_text.insert(tk.END, f"Excel Spreadsheet: {os.path.basename(excel_path)}\n\n")
                
                if validation_result["saved_mapping_used"]:
                    validation_text.insert(tk.END, "✅ Saved mapping loaded successfully\n\n", "success")
                
                if validation_result["missing"]:
                    validation_text.insert(tk.END, "❌ Fields missing in Excel:\n", "error")
                    for field in validation_result["missing"]:
                        validation_text.insert(tk.END, f"• {field}\n")
                    validation_text.insert(tk.END, "\n")
                
                if validation_result["warnings"]:
                    validation_text.insert(tk.END, "⚠️ Warnings:\n", "warning")
                    for warning in validation_result["warnings"]:
                        validation_text.insert(tk.END, f"• {warning}\n")
                    validation_text.insert(tk.END, "\n")
                
                if validation_result["suggestions"]:
                    validation_text.insert(tk.END, "💡 Similar column suggestions:\n", "info")
                    for field, matches in validation_result["suggestions"].items():
                        validation_text.insert(tk.END, f"• For '{field}': {', '.join(matches)}\n")
                    validation_text.insert(tk.END, "\n")
                
                validation_text.insert(tk.END, "📊 Columns available in Excel:\n")
                for col in validation_result["columns"]:
                    validation_text.insert(tk.END, f"• {col}\n")
                
                validation_text.config(state=tk.DISABLED)
            
            btn_frame = ttk.Frame(fields_window)
            btn_frame.pack(fill=tk.X, pady=10)
            
            ttk.Button(btn_frame, text="Save Mapping", 
                      command=lambda: self.save_column_mapping(template_path)).pack(side=tk.LEFT)
            ttk.Button(btn_frame, text="Close", command=fields_window.destroy).pack(side=tk.RIGHT)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to detect fields: {str(e)}")
            logger.exception("Error detecting template fields")
        finally:
            self.show_indeterminate_progress(False)
            self.update_status("Ready")
    
    def save_column_mapping(self, template_path: str):
        """Saves current column mapping for the template"""
        if not template_path:
            messagebox.showwarning("Warning", "No template selected")
            return
        
        excel_path = self.excel_path.get()
        if not excel_path or not os.path.exists(excel_path):
            messagebox.showwarning("Warning", "No Excel spreadsheet selected")
            return
        
        try:
            
            _, validation_result = self.declaration_gen.validate_columns(excel_path, template_path, self.config)
            
            if not validation_result.get("mapping"):
                messagebox.showwarning("Warning", "No mapping to save")
                return
            
            self.config.save_column_mapping(template_path, validation_result["mapping"])
            messagebox.showinfo("Success", "Column mapping saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save mapping: {str(e)}")
            logger.exception("Error saving column mapping")
    
    def show_preview(self):
        template_path = self.template_path.get()
        if not template_path:
            messagebox.showwarning("Warning", "No template selected")
            return
        
        self.show_indeterminate_progress(True)
        self.update_status("Preparing preview...")
        
        try:
            preview_window = tk.Toplevel(self.root)
            preview_window.title("Document Preview")
            preview_window.geometry("900x700")
            
            notebook = ttk.Notebook(preview_window)
            notebook.pack(fill=tk.BOTH, expand=True)
            
            doc_tab = ttk.Frame(notebook)
            notebook.add(doc_tab, text="Preview")
            
            canvas = tk.Canvas(doc_tab)
            scrollbar = ttk.Scrollbar(doc_tab, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(
                    scrollregion=canvas.bbox("all")
                )
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            example_data = {
                "{{NOME}}": "Erlon Lopes",
                "{{CPF}}": "123.456.789-00",
                "{{MATRICULA}}": "20230001",
                "{{EMAIL}}": "erlon.araujo@upe.br",
                "{{TIPO_DOC}}": "delaração ou defesa", 
                "{{TITULO}}": "The goal is simple: improve every day.",
                "{{DATA_DEFESA}}": "23/09/2023",
                "{{LOCAL}}": "Remote - Google Meet",
                "{{HORARIO}}": "14",
                "{{NOME_ORIENTADOR}}": "A Person",
                "{{NOME_EXAMINADOR1}}": "Another Person",
                "{{NOME_EXAMINADOR2}}": "One More Person",
                "{{DATA}}": datetime.now().strftime("%d/%m/%Y"),
                "{{DIA}}": str(datetime.now().day),
                "{{MES}}": str(datetime.now().month),
                "{{ANO}}": str(datetime.now().year),
            }
            
            doc = DocxDocument(template_path)
            
            ttk.Label(scrollable_frame, 
                     text="Document Preview", 
                     font=('Arial', 12, 'bold')).pack(pady=10)
            
            for para in doc.paragraphs:
                text = para.text
                if text.strip():
                    for key, value in example_data.items():
                        text = text.replace(key, f"[{value}]")
                    
                    frame = ttk.Frame(scrollable_frame)
                    frame.pack(fill=tk.X, padx=10, pady=2)
                    
                    ttk.Label(frame, text=para.style.name, width=15, 
                             font=('Arial', 8), foreground="gray").pack(side=tk.LEFT)
                    
                    text_label = ttk.Label(frame, text=text, wraplength=600, justify=tk.LEFT)
                    text_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            info_tab = ttk.Frame(notebook)
            notebook.add(info_tab, text="Informação")
            
            ttk.Label(info_tab, text="Este é um exemplo de substituição de campo:", 
                     font=('Arial', 10, 'bold')).pack(pady=10)
            
            for key, value in example_data.items():
                frame = ttk.Frame(info_tab)
                frame.pack(fill=tk.X, padx=10, pady=2)
                ttk.Label(frame, text=key, width=20, anchor=tk.W).pack(side=tk.LEFT)
                ttk.Label(frame, text=value, foreground="blue").pack(side=tk.LEFT)
            
            ttk.Button(preview_window, text="Close", command=preview_window.destroy).pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate preview: {str(e)}")
            logger.exception("Error generating preview")
        finally:
            self.show_indeterminate_progress(False)
            self.update_status("Ready")
    
    def validate_paths(self, *entries) -> bool:
        missing = []
        for entry, name in entries:
            if isinstance(entry, tk.Entry):
                value = entry.get().strip()
            elif isinstance(entry, tk.Listbox):
                value = entry.size()
            else:
                value = str(entry).strip()
            
            if not value:
                missing.append(name)
        
        if missing:
            message = "Required fields missing:\n- " + "\n- ".join(missing)
            messagebox.showerror("Error", message)
            return False
        return True
        
    def start_declaration_generation(self):
        if not self.validate_paths(
            (self.excel_path, "Planilha Excel"),
            (self.template_path, "Modelo do Word"),
            (self.output_folder, "Pasta de saída")
        ):
            return
        
        if self.declaration_gen.running:
            messagebox.showwarning("Aviso”, “Um processo já está em execução")
            return
        
        template_path = self.template_path.get()
        excel_path = self.excel_path.get()
        
        self.show_indeterminate_progress(True)
        self.update_status("Detectando campos de modelo...")
        
        placeholders = self.declaration_gen.detect_placeholders(template_path)
        if not placeholders:
            messagebox.showwarning("Aviso: Nenhum campo detectado no modelo. Verifique o modelo.")
            self.show_indeterminate_progress(False)
            self.update_status("Ready")
            return
        
        self.update_status("Validating Excel columns...")
        
        valid, validation_result = self.declaration_gen.validate_columns(excel_path, template_path, self.config)
        if not valid and validation_result.get("missing"):
            msg = "Required columns missing in Excel:\n- " + "\n- ".join(validation_result["missing"])
            if validation_result.get("suggestions"):
                msg += "\n\nSimilar column suggestions:\n"
                for field, matches in validation_result["suggestions"].items():
                    msg += f"- For '{field}': {', '.join(matches)}\n"
            
            if not messagebox.askyesno("Warning", msg + "\nContinue anyway?"):
                self.show_indeterminate_progress(False)
                self.update_status("Ready")
                return
        
        self.show_indeterminate_progress(False)
        self.clear_log(self.declaration_log)

        # Configura o gerador com as opções da interface
        self.declaration_gen.save_in_subfolders = self.save_in_subfolders_var.get()
        if hasattr(self.declaration_gen, 'custom_filename_pattern'):
            self.declaration_gen.custom_filename_pattern = self.filename_pattern_entry.get().strip()
        
        thread = threading.Thread(
            target=self.declaration_gen.generate_declarations,
            args=(
                self.excel_path.get(),
                self.template_path.get(),
                self.output_folder.get(),
                self.zip_output.get(),
                self.config,
                self.filename_pattern_entry.get().strip()  # Passa o padrão explicitamente
            ),
            daemon=True
        )
        thread.start()


    def convert_docx_to_pdf(self):
        if not self.validate_paths(
            (self.docx_listbox, "DOCX Files"),
            (self.docx_output_folder, "Output Folder")
        ):
            return
        
        if self.docx_converter.running:
            messagebox.showwarning("Aviso", "Um processo já está em execução")
            return
        
        docx_paths = list(self.docx_listbox.get(0, tk.END))
        output_folder = self.docx_output_folder.get()
        
        self.clear_log(self.docx_log)
        
        self.current_processor = self.docx_converter
        thread = threading.Thread(
            target=self.docx_converter.convert_docx_to_pdf,
            args=(docx_paths, output_folder),
            daemon=True
        )
        thread.start()
    
    def convert_pdf_to_docx(self):
        if not self.validate_paths(
            (self.pdf_listbox, "PDF Files"),
            (self.pdf_output_folder, "Output Folder")
        ):
            return
        
        if self.pdf_converter.running:
            messagebox.showwarning("Aviso", "Um processo já está em execução")
            return
        
        pdf_paths = list(self.pdf_listbox.get(0, tk.END))
        output_folder = self.pdf_output_folder.get()
        
        self.clear_log(self.pdf_log)
        
        self.current_processor = self.pdf_converter
        thread = threading.Thread(
            target=self.pdf_converter.convert_pdf_to_docx,
            args=(pdf_paths, output_folder),
            daemon=True
        )
        thread.start()
    
    def stop_process(self):
        if self.current_processor and self.current_processor.running:
            if messagebox.askyesno("Confirmar", "Deseja realmente parar de processar?"):
                self.current_processor.stop()
                self.update_status("Parando o processo...(Pode demorar um pouco)")
        else:
            messagebox.showinfo("Informações", "Nenhum processo em execução para parar")
    
    def process_queues(self):
        try:
            if self.current_processor:
                # Process log updates
                while not self.current_processor.log_queue.empty():
                    message, level = self.current_processor.log_queue.get()
                    log_widget = self.get_current_log_widget()
                    if log_widget:
                        self.log_message(log_widget, message, level)

                # Process errors
                while not self.current_processor.error_queue.empty():
                    args, kwargs, error, trace = self.current_processor.error_queue.get()
                    logger.error(f"Detailed error: {error}\nArgs: {args}\nTrace: {trace}")

                # Process progress
                while not self.current_processor.progress_queue.empty():
                    value, max_value = self.current_processor.progress_queue.get()
                    if max_value is not None:
                        self.progress["maximum"] = max_value
                    self.progress["value"] = value

                    if max_value and max_value > 0:
                        percent = int((value / max_value) * 100)
                        self.progress_label.config(text=f"{percent}%")

                
                while not self.current_processor.gui_update_queue.empty():
                    widget, method, args, kwargs = self.current_processor.gui_update_queue.get()
                    if hasattr(widget, method):
                        try:
                            getattr(widget, method)(*args, **kwargs)
                        except Exception as e:
                            logger.error(f"GUI update error: {str(e)}")

            
                if not self.current_processor.running and self.current_processor.end_time:
               
                    report = self.current_processor._generate_execution_report()
                    log_widget = self.get_current_log_widget()
                    log_text = log_widget.get("1.0", tk.END) if log_widget else ""

                    if self.current_processor == self.declaration_gen:
                        input_files = f"Excel: {self.excel_path.get()}\nTemplate: {self.template_path.get()}"
                        output_folder = self.output_folder.get()
                        operation_type = "Declaration Generation"
                    elif self.current_processor == self.docx_converter:
                        input_files = "\n".join(self.docx_listbox.get(0, tk.END))
                        output_folder = self.docx_output_folder.get()
                        operation_type = "DOCX to PDF"
                    elif self.current_processor == self.pdf_converter:
                        input_files = "\n".join(self.pdf_listbox.get(0, tk.END))
                        output_folder = self.pdf_output_folder.get()
                        operation_type = "PDF to DOCX"
                    else:
                        input_files = "Unknown"
                        output_folder = "Unknown"
                        operation_type = "Unknown"

                    self.config.history.add_record(
                        operation_type=operation_type,
                        input_files=input_files,
                        output_folder=output_folder,
                        status="completed" if not self.current_processor.stop_requested else "stopped",
                        processed_count=self.current_processor.processed_count,
                        error_count=self.current_processor.error_count,
                        duration=report["duration"],
                        log_text=log_text
                    )

                  
                    if self.config.config.get("show_notifications", True):
                        self.show_notification(
                            "Processing Complete",
                            f"{operation_type} finished\n"
                            f"Processed: {self.current_processor.processed_count}\n"
                            f"Errors: {self.current_processor.error_count}"
                        )

                    
                    if (self.config.config.get("enable_sounds", True)
                            and platform.system() == "Windows"):
                        winsound.MessageBeep()

                    self.current_processor = None
                    self.update_status("Ready")
                    self.refresh_history()

            self.root.after(100, self.process_queues)

        except Exception as e:
            logger.error(f"Error in process_queues: {str(e)}")
            self.root.after(100, self.process_queues)

    def show_notification(self, title: str, message: str):

        if self.notification_window and self.notification_window.winfo_exists():
            self.notification_window.destroy()
        
        self.notification_window = tk.Toplevel(self.root)
        self.notification_window.title(title)
        self.notification_window.geometry("300x150")
        self.notification_window.resizable(False, False)
        
        # Position near system tray
        screen_width = self.notification_window.winfo_screenwidth()
        screen_height = self.notification_window.winfo_screenheight()
        x = screen_width - 320
        y = screen_height - 200
        self.notification_window.geometry(f"+{x}+{y}")
        
        ttk.Label(self.notification_window, text=message, wraplength=280).pack(pady=20)
        ttk.Button(self.notification_window, text="OK", command=self.notification_window.destroy).pack(pady=10)
        
        
        self.notification_window.after(5000, self.notification_window.destroy)
    
    def get_current_log_widget(self):
        if self.current_processor == self.declaration_gen:
            return self.declaration_log
        elif self.current_processor == self.docx_converter:
            return self.docx_log
        elif self.current_processor == self.pdf_converter:
            return self.pdf_log
        return None
    
    def log_message(self, log_widget: ScrolledText, message: str, level: str = "info"):
        if not log_widget:
            return
            
        log_widget.config(state=tk.NORMAL)
        
        if level == "error":
            log_widget.insert(tk.END, message + "\n", "error")
        elif level == "warning":
            log_widget.insert(tk.END, message + "\n", "warning")
        elif level == "success":
            log_widget.insert(tk.END, message + "\n", "success")
        elif level == "info":
            log_widget.insert(tk.END, message + "\n", "info")
        else:
            log_widget.insert(tk.END, message + "\n")
        
        log_widget.see(tk.END)
        log_widget.config(state=tk.DISABLED)
    
    def clear_log(self, log_widget: ScrolledText):
        log_widget.config(state=tk.NORMAL)
        log_widget.delete(1.0, tk.END)
        log_widget.config(state=tk.DISABLED)
    
    def update_status(self, message: str):
        self.status_label.config(text=message)
    
    def on_closing(self):
        self.save_config()
        if self.current_processor and self.current_processor.running:
            if messagebox.askokcancel("Exit", "A process is running. Really exit?"):
                self.current_processor.stop()
                self.root.destroy()
        else:
            self.root.destroy()

    def show_multi_model_dialog(self):
        """Mostra o diálogo para agendamento de múltiplos modelos"""
        if not hasattr(self, 'task_manager'):
            messagebox.showerror("Erro", "TaskManager não inicializado")
            return
        
        #  template_manager como parâmetro adicional
        MultiModelDialog(self.root, self.task_manager, self.template_manager)

class TaskManagerDeclaracoes:
    def __init__(self, scheduler, config, dialog, max_workers=4, use_libreoffice=False, libreoffice_path=""):
        self.scheduler = scheduler
        self.config = config
        self.dialog = dialog
        self.max_workers = max_workers
        self.use_libreoffice = use_libreoffice
        self.libreoffice_path = libreoffice_path

    def agendar_multiplos_modelos(self, excel_path: str, modelos: List[dict], pasta_saida: str):
        """Agenda múltiplos modelos para processamento."""
        for modelo in modelos:
            self.agendar_declaracao(
                excel_path=excel_path,
                modelo_path=modelo['modelo_path'],
                pasta_saida=pasta_saida,
                filename_pattern=modelo.get('filename_pattern', "{{Nome}}_{{TipoDoc}}"),
                save_in_subfolders=modelo.get('save_in_subfolders', False)
            )

    def agendar_declaracao(self, excel_path: str, modelo_path: str, pasta_saida: str, 
                        criar_zip: bool = False, filename_pattern: str = None,
                        save_in_subfolders: bool = False):
        """Agenda com todos os parâmetros necessários"""
        task_args = {
            'excel_path': excel_path,
            'modelo_path': modelo_path,
            'pasta_saida': pasta_saida,
            'criar_zip': criar_zip,
            'filename_pattern': filename_pattern or "{{Nome}}_{{TipoDoc}}",
            'save_in_subfolders': save_in_subfolders
        }
        
        self.scheduler.schedule_task(
            task_func=self._executar_geracao_com_config,
            task_args=(task_args,),  # Note a vírgula para criar uma tupla
            run_time=datetime.now(),
            description=f"Geração: {os.path.basename(modelo_path)}",
            dialog=self.dialog
        )

    def _executar_geracao_com_config(self, task_args):
        try:
            generator = BatchDeclarationGenerator(
                max_workers=self.max_workers,
                use_libreoffice=self.use_libreoffice,
                libreoffice_path=self.libreoffice_path
            )
            
            generator.save_in_subfolders = task_args.get('save_in_subfolders', False)
            if hasattr(generator, 'custom_filename_pattern'):
                generator.custom_filename_pattern = task_args.get('filename_pattern', "{{Nome}}_{{TipoDoc}}")
            
            # Adiciona logging para diagnóstico
            logger.info(f"Iniciando geração para modelo: {task_args['modelo_path']}")
            
            generator.generate_declarations(
                excel_path=task_args['excel_path'],
                template_path=task_args['modelo_path'],
                output_folder=task_args['pasta_saida'],
                create_zip=task_args.get('criar_zip', False),
                config=self.config,
                nome_saida_template=task_args.get('filename_pattern', "{{Nome}}_{{TipoDoc}}")
            )
            
            logger.info(f"Geração concluída para: {task_args['modelo_path']}")
            
        except Exception as e:
            logger.error(f"Falha na geração do documento: {str(e)}")
            # Adicione aqui qualquer tratamento de erro adicional que desejar
            raise  # Re-lança a exceção para o scheduler
    def cancelar_tarefas(self):
        """Cancela todas as tarefas pendentes"""
        self.scheduler.task_queue.clear()
        for task_id in list(self.scheduler.task_log.keys()):
            self.scheduler.task_log[task_id]["status"] = "cancelled"

    def _executar_geracao_com_nome(self, task_args):
        """Executa a geração aplicando o padrão de nomes"""
        try:
            generator = BatchDeclarationGenerator(
                max_workers=self.max_workers,
                use_libreoffice=self.use_libreoffice,
                libreoffice_path=self.libreoffice_path
            )
            
            # Configuração explícita do padrão de nomes
            generator.custom_filename_pattern = task_args['filename_pattern']
            generator.save_in_subfolders = task_args['save_in_subfolders']
            generator.subfolder_column = "Nome"  # Ou outra coluna relevante

            # Validação e execução
            generator.generate_declarations(
                excel_path=task_args['excel_path'],
                template_path=task_args['modelo_path'],
                output_folder=task_args['pasta_saida'],
                create_zip=task_args['criar_zip'],
                config=task_args['config'],
                nome_saida_template=task_args['filename_pattern']  # Passa o padrão aqui
            )
        except Exception as e:
            logging.error(f"Falha na tarefa agendada: {str(e)}")

class TemplateManager:
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = Path(templates_dir)
        self.config_path = self.templates_dir / "_config.json"
        self.templates = self._carregar_config()

    def _carregar_config(self) -> Dict:
        """Carrega a configuração dos templates"""
        if not self.config_path.exists():
            return {"templates": {}}
        
        with open(self.config_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _salvar_config(self):
        """Salva as alterações no arquivo de configuração"""
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(self.templates, f, indent=4, ensure_ascii=False)

    def adicionar_template(
        self,
        nome: str,
        arquivo: str,
        categoria: str,
        tags: List[str],
        campos_obrigatorios: List[str],
        descricao: str = ""
    ):
        """Adiciona um novo template à biblioteca"""
        rel_path = str(Path(arquivo).relative_to(self.templates_dir))
        
        self.templates["templates"][nome] = {
            "categoria": categoria,
            "arquivo": rel_path,
            "tags": tags,
            "campos_obrigatorios": campos_obrigatorios,
            "descricao": descricao
        }
        self._salvar_config()

    def buscar_templates(
        self,
        categoria: str = None,
        tag: str = None,
        termo: str = None
    ) -> List[Dict]:
        """Busca templates com filtros"""
        resultados = []
        
        for nome, template in self.templates["templates"].items():
            if categoria and template["categoria"] != categoria:
                continue
                
            if tag and tag not in template["tags"]:
                continue
                
            if termo and termo.lower() not in nome.lower():
                continue
                
            resultados.append({
                "nome": nome,
                **template
            })
            
        return resultados

    def get_caminho_template(self, nome_template: str) -> Path:
        """Retorna o caminho absoluto do arquivo do template"""
        if nome_template not in self.templates["templates"]:
            raise ValueError(f"Template '{nome_template}' não encontrado")
            
        return self.templates_dir / self.templates["templates"][nome_template]["arquivo"]

  
    def validar_dados(self, template_name: str, dados: dict) -> tuple:
        """
        Valida se os dados contêm todos os campos obrigatórios
        Retorna: (bool, lista_de_campos_faltantes)
        """
        template = self.templates["templates"].get(template_name)
        if not template:
            return False, ["Template não encontrado"]
        
        faltantes = [
            campo for campo in template["campos_obrigatorios"]
            if campo not in dados
        ]
        
        return (not faltantes), faltantes

    def _salvar_config(self):
        """Salva com backup automático"""
        TemplateBackup(self).create_backup()  # Cria backup antes de salvar
        with open(self.config_path, 'w') as f:
            json.dump(self.templates, f, indent=4)

class TemplateLibraryDialog(tk.Toplevel):
    def __init__(self, parent, template_manager, multi_select=False):
        super().__init__(parent)
        self.title("Biblioteca de Templates")
        self.template_manager = template_manager
        self.multi_select = multi_select
        self.selected_template = None
        self.selected_templates = []
        self.auto_refresh = tk.BooleanVar(value=True)
            
        # Initialize auto_refresh before it's used
        self.auto_refresh = tk.BooleanVar(value=True)
        
        # Configuração da janela
        self.geometry("1000x800")
        self.resizable(True, True)
        
        # Frame principal
        main_frame = ttk.Frame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Filtros
        self._setup_filters(main_frame)
        
        # Lista de templates
        self._setup_template_list(main_frame)
        
        # Detalhes
        self._setup_details_panel(main_frame)
        
        # Botões
        self._setup_buttons(main_frame)
        
        # Carrega os templates
        self._load_templates()
        
        # Atualização automática
        self.after(30000, self._auto_refresh_list)
        
        # Foca na janela
        self.grab_set()
        self.focus_set()

    def _setup_buttons(self, parent):
        """Configura os botões de ação"""
        button_frame = ttk.Frame(parent)
        button_frame.pack(fill=tk.X, pady=10)
        
        if self.multi_select:
            ttk.Button(
                button_frame, 
                text="Selecionar Templates", 
                command=self._on_select_multiple
            ).pack(side=tk.LEFT, padx=5)
        else:
            ttk.Button(
                button_frame, 
                text="Selecionar Template", 
                command=self._on_select_single
            ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame, 
            text="Cancelar", 
            command=self.destroy
        ).pack(side=tk.RIGHT, padx=5)
        
        ttk.Checkbutton(
            button_frame, 
            text="Atualização Automática", 
            variable=self.auto_refresh
        ).pack(side=tk.LEFT, padx=10)
    
    def _setup_filters(self, parent):
        """Configura os controles de filtro"""
        filter_frame = ttk.LabelFrame(parent, text="Filtros", padding=10)
        filter_frame.pack(fill=tk.X, pady=5)
        
        # Filtro por categoria
        ttk.Label(filter_frame, text="Categoria:").pack(side=tk.LEFT)
        self.categoria_var = tk.StringVar()
        categorias = ["Todos"] + self._get_categorias()
        ttk.OptionMenu(filter_frame, self.categoria_var, "Todos", *categorias, 
                      command=lambda _: self._load_templates()).pack(side=tk.LEFT, padx=5)
        
        # Filtro por tag
        ttk.Label(filter_frame, text="Tag:").pack(side=tk.LEFT)
        self.tag_var = tk.StringVar()
        tags = ["Todas"] + self._get_tags()
        ttk.OptionMenu(filter_frame, self.tag_var, "Todas", *tags, 
                      command=lambda _: self._load_templates()).pack(side=tk.LEFT, padx=5)
        
        # Campo de busca
        ttk.Label(filter_frame, text="Buscar:").pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(filter_frame, textvariable=self.search_var)
        search_entry.pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        search_entry.bind("<KeyRelease>", lambda _: self._load_templates())
    
    def _setup_template_list(self, parent):
        """Configura a lista de templates"""
        list_frame = ttk.LabelFrame(parent, text="Templates Disponíveis", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Treeview com colunas
        self.tree = ttk.Treeview(list_frame, columns=("name", "category", "tags", "required"), 
                               selectmode="extended" if self.multi_select else "browse")
        
        # Configuração das colunas
        self.tree.heading("#0", text="ID", anchor=tk.W)
        self.tree.heading("name", text="Nome", anchor=tk.W, 
                         command=lambda: self._sort_column("name", False))
        self.tree.heading("category", text="Categoria", anchor=tk.W,
                         command=lambda: self._sort_column("category", False))
        self.tree.heading("tags", text="Tags", anchor=tk.W,
                         command=lambda: self._sort_column("tags", False))
        self.tree.heading("required", text="Campos Obrig.", anchor=tk.W,
                         command=lambda: self._sort_column("required", False))
        
        self.tree.column("#0", width=50, stretch=tk.NO)
        self.tree.column("name", width=200, stretch=tk.YES)
        self.tree.column("category", width=150, stretch=tk.YES)
        self.tree.column("tags", width=150, stretch=tk.YES)
        self.tree.column("required", width=150, stretch=tk.YES)
        
        # Scrollbars
        vsb = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(list_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        # Layout
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        
        # Configura expansão
        list_frame.grid_rowconfigure(0, weight=1)
        list_frame.grid_columnconfigure(0, weight=1)
        
        # Bindings
        self.tree.bind("<Double-1>", self._on_double_click)
        self.tree.bind("<<TreeviewSelect>>", self._show_template_details)
    
    def _setup_details_panel(self, parent):
        """Configura o painel de detalhes"""
        details_frame = ttk.LabelFrame(parent, text="Detalhes do Template", padding=10)
        details_frame.pack(fill=tk.BOTH, pady=5)
        
        # Texto com scroll
        self.details_text = ScrolledText(details_frame, height=8, wrap=tk.WORD, state=tk.DISABLED)
        self.details_text.pack(fill=tk.BOTH, expand=True)
        
        # Frame para informações adicionais
        info_frame = ttk.Frame(details_frame)
        info_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(info_frame, text="Caminho:").pack(side=tk.LEFT)
        self.path_label = ttk.Label(info_frame, text="", foreground="blue")
        self.path_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
    
    def _setup_buttons(self, parent):
        """Configura os botões de ação"""
        btn_frame = ttk.Frame(parent)
        btn_frame.pack(fill=tk.X, pady=10)
        
        # Botões de seleção
        if self.multi_select:
            ttk.Button(btn_frame, text="Selecionar Tudo", 
                      command=self._select_all).pack(side=tk.LEFT)
            ttk.Button(btn_frame, text="Limpar Seleção", 
                      command=self._clear_selection).pack(side=tk.LEFT)
        
        # Botões de ação
        ttk.Button(btn_frame, text="Pré-visualizar", 
                  command=self._preview_template).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Validar", 
                  command=self._validate_template).pack(side=tk.LEFT, padx=5)
        
        # Checkbox para atualização automática
        ttk.Checkbutton(btn_frame, text="Atualização automática", 
                       variable=self.auto_refresh).pack(side=tk.LEFT, padx=10)
        
        # Botões de finalização
        ttk.Button(btn_frame, text="Cancelar", 
                  command=self.destroy).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Selecionar", 
                  command=self._finalize_selection).pack(side=tk.RIGHT, padx=5)
    
    def _load_templates(self):
        """Carrega os templates na lista"""
        try:
            # Limpa a lista atual
            for item in self.tree.get_children():
                self.tree.delete(item)
            
            # Obtém os filtros
            categoria = self.categoria_var.get()
            tag = self.tag_var.get()
            search_term = self.search_var.get().lower()
            
            # Carrega os templates filtrados
            templates = self.template_manager.buscar_templates(
                categoria=None if categoria == "Todos" else categoria,
                tag=None if tag == "Todas" else tag
            )
            
            # Adiciona à treeview
            for idx, template in enumerate(templates):
                self.tree.insert("", "end", iid=str(idx), 
                            values=(
                                template["nome"],
                                template["categoria"],
                                ", ".join(template["tags"]),
                                ", ".join(template["campos_obrigatorios"])
                            ))
        except Exception as e:
            print(f"Erro ao carregar templates: {str(e)}")
        
    def _sort_column(self, col, reverse):
        """Ordena a treeview pela coluna clicada"""
        data = [(self.tree.set(child, col), child) for child in self.tree.get_children("")]
        data.sort(reverse=reverse)
        
        for index, (val, child) in enumerate(data):
            self.tree.move(child, "", index)
        
        self.tree.heading(col, command=lambda: self._sort_column(col, not reverse))
    
    def _show_template_details(self, template_name):
        """Mostra os detalhes do template selecionado"""
        try:
            # Corrigindo o acesso aos dados do template
            template = self.template_manager.templates["templates"].get(template_name)
            if not template:
                return
                
            # Limpa os detalhes anteriores
            self.details_text.config(state=tk.NORMAL)
            self.details_text.delete(1.0, tk.END)
            
            # Exibe os detalhes
            self.details_text.insert(tk.END, f"Nome: {template_name}\n")
            self.details_text.insert(tk.END, f"Categoria: {template.get('categoria', 'N/A')}\n")
            self.details_text.insert(tk.END, f"Tags: {', '.join(template.get('tags', []))}\n")
            self.details_text.insert(tk.END, f"Campos obrigatórios: {', '.join(template.get('campos_obrigatorios', []))}\n")
            self.details_text.insert(tk.END, f"\nDescrição:\n{template.get('descricao', 'Nenhuma descrição disponível')}")
            
            self.details_text.config(state=tk.DISABLED)
        except Exception as e:
            print(f"Erro ao mostrar detalhes: {str(e)}")
        
    def _preview_template(self):
        """Mostra uma prévia do template selecionado"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Aviso", "Nenhum template selecionado")
            return
        
        template_name = self.tree.item(selected[0])["values"][0]
        try:
            template_path = self.template_manager.get_caminho_template(template_name)
            
            # Tenta abrir com aplicativo padrão
            if platform.system() == "Windows":
                os.startfile(template_path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", template_path])
            else:  # linux
                subprocess.run(["xdg-open", template_path])
                
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir o template:\n{str(e)}")
    
    def _validate_template(self):
        """Valida o template selecionado"""
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Aviso", "Nenhum template selecionado")
            return
        
        template_name = self.tree.item(selected[0])["values"][0]
        template = next((t for t in self.template_manager.templates["templates"].values() 
                        if t["nome"] == template_name), None)
        
        if template:
            try:
                campos_faltantes = []
                doc = DocxDocument(self.template_manager.get_caminho_template(template_name))
                
                # Verifica parágrafos
                for campo in template["campos_obrigatorios"]:
                    placeholder = f"{{{{{campo}}}}}"
                    found = False
                    
                    for para in doc.paragraphs:
                        if placeholder in para.text:
                            found = True
                            break
                    
                    # Verifica tabelas se não encontrou nos parágrafos
                    if not found:
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if placeholder in cell.text:
                                        found = True
                                        break
                                if found:
                                    break
                            if found:
                                break
                    
                    if not found:
                        campos_faltantes.append(campo)
                
                if campos_faltantes:
                    messagebox.showwarning("Validação", 
                                         f"Campos obrigatórios faltando no template:\n\n" + 
                                         "\n".join(f"• {campo}" for campo in campos_faltantes))
                else:
                    messagebox.showinfo("Validação", "✓ Todos os campos obrigatórios estão presentes no template!")
            
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao validar template:\n{str(e)}")
    
    def _on_double_click(self, event):
        """Trata duplo clique na lista"""
        if not self.multi_select:
            self._finalize_selection()
    
    def _finalize_selection(self):
        """Finaliza a seleção e fecha o diálogo"""
        if self.multi_select:
            self.selected_templates = [
                self.tree.item(item)["values"][0]  # Pega o nome do template
                for item in self.tree.selection()
            ]
            if not self.selected_templates:
                messagebox.showwarning("Aviso", "Nenhum template selecionado")
                return
        else:
            selected = self.tree.selection()
            if not selected:
                messagebox.showwarning("Aviso", "Nenhum template selecionado")
                return
            self.selected_template = self.tree.item(selected[0])["values"][0]
        
        self.destroy()
    
    def _select_all(self):
        """Seleciona todos os templates na lista"""
        for item in self.tree.get_children():
            self.tree.selection_add(item)
    
    def _clear_selection(self):
        """Limpa a seleção atual"""
        self.tree.selection_set([])
    
    def _auto_refresh_list(self):
        """Atualiza a lista automaticamente se habilitado"""
        if self.auto_refresh.get():
            self._load_templates()
        self.after(30000, self._auto_refresh_list)
    
    def _get_categorias(self):
        """Retorna lista de categorias disponíveis"""
        return list(set(
            t["categoria"] 
            for t in self.template_manager.templates["templates"].values()
        ))
    
    def _get_tags(self):
        """Retorna lista de tags disponíveis"""
        tags = set()
        for t in self.template_manager.templates["templates"].values():
            tags.update(t["tags"])
        return sorted(tags)

class TemplateBatchImporter:
    def __init__(self, template_manager: TemplateManager):
        self.tm = template_manager

    def import_from_folder(self, folder_path: str):
        """Importa todos os .docx de uma pasta"""
        for docx_path in Path(folder_path).glob("*.docx"):
            self.tm.adicionar_template(
                nome=f"batch_{docx_path.stem}",
                arquivo=str(docx_path),
                categoria="importados",
                tags=["batch_import"],
                campos_obrigatorios=self._detect_fields(docx_path),
                descricao=f"Importado em {datetime.now():%d/%m/%Y}"
            )

    def _detect_fields(self, file_path) -> list:
        """Detecta campos automaticamente (usando regex no texto)"""
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return list(set(re.findall(r'\{\{(\w+)\}\}', text)))

class MultiModelDialog(tk.Toplevel):
    def __init__(self, parent, task_manager, template_manager):
        super().__init__(parent)
        self.title("Agendar Múltiplos Modelos")
        self.task_manager = task_manager
        self.template_manager = template_manager
        
        # Configuração da planilha
        ttk.Label(self, text="Planilha Excel:").grid(row=0, column=0, sticky=tk.W)
        self.excel_path = ttk.Entry(self, width=40)
        self.excel_path.grid(row=0, column=1)
        ttk.Button(self, text="Procurar", command=self.select_excel).grid(row=0, column=2)
        
        # Configuração da pasta de saída
        ttk.Label(self, text="Pasta de Saída:").grid(row=1, column=0, sticky=tk.W)
        self.output_folder = ttk.Entry(self, width=40)
        self.output_folder.grid(row=1, column=1)
        ttk.Button(self, text="Procurar", command=self.select_output).grid(row=1, column=2)
        
        # Lista de modelos
        ttk.Label(self, text="Modelos:").grid(row=2, column=0, sticky=tk.W)
        self.models_frame = ttk.Frame(self)
        self.models_frame.grid(row=3, column=0, columnspan=3, sticky=tk.EW)
        
        self.model_entries = []
        self.add_model_row()
        
        ttk.Button(self, text="+ Adicionar Modelo", command=self.add_model_row).grid(row=4, column=0)
        
        # Botão para selecionar da biblioteca
        ttk.Button(self, text="Selecionar da Biblioteca", 
                  command=self.select_from_library).grid(row=4, column=1, sticky=tk.W)
        
        # Botões de ação
        btn_frame = ttk.Frame(self)
        btn_frame.grid(row=5, column=0, columnspan=3, pady=10)
        ttk.Button(btn_frame, text="Cancelar", command=self.destroy).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="Gerar", command=self.schedule_tasks).pack(side=tk.RIGHT)
    
    def select_from_library(self):
        """Abre a biblioteca para selecionar templates"""
        def on_select(template_names):
            if not template_names:
                return
                
            # Adiciona uma linha para cada template selecionado
            for template_name in template_names:
                self.add_model_row()  # Adiciona nova linha
                last_entry = self.model_entries[-1]  # Pega a última linha adicionada
                template_path = self.template_manager.get_caminho_template(template_name)
                last_entry['modelo_path'].delete(0, tk.END)
                last_entry['modelo_path'].insert(0, str(template_path))
        
        # Cria e mostra o diálogo da biblioteca com seleção múltipla
        library_dialog = TemplateLibraryDialog(self, self.template_manager, multi_select=True)
        self.wait_window(library_dialog)
        
        if library_dialog.selected_templates:
            on_select(library_dialog.selected_templates)
    
    def add_model_row(self):
        row = len(self.model_entries)
        
        # Modelo
        ttk.Label(self.models_frame, text=f"Modelo {row+1}:").grid(row=row, column=0)
        modelo_path = ttk.Entry(self.models_frame, width=30)
        modelo_path.grid(row=row, column=1)
        ttk.Button(self.models_frame, text="Procurar", 
                  command=lambda: self.select_model(modelo_path)).grid(row=row, column=2)
        
        # Padrão de nome
        ttk.Label(self.models_frame, text="Padrão do Nome:").grid(row=row, column=3)
        filename_pattern = ttk.Entry(self.models_frame, width=30)
        filename_pattern.insert(0, "{{Nome}}_{{TipoDoc}}.docx")
        filename_pattern.grid(row=row, column=4)
        
        # Subpastas
        subfolder_var = tk.BooleanVar()
        ttk.Checkbutton(self.models_frame, text="Salvar em subpastas", 
                       variable=subfolder_var).grid(row=row, column=5)
        
        self.model_entries.append({
            'modelo_path': modelo_path,
            'filename_pattern': filename_pattern,
            'subfolder_var': subfolder_var
        })
    
    def select_excel(self):
        path = filedialog.askopenfilename(
            title="Selecione a planilha Excel",
            filetypes=[("Excel Files", "*.xlsx;*.xls"), ("All files", "*.*")]
        )
        if path:
            self.excel_path.delete(0, tk.END)
            self.excel_path.insert(0, path)
    
    def select_output(self):
        path = filedialog.askdirectory(title="Selecione a pasta de saída")
        if path:
            self.output_folder.delete(0, tk.END)
            self.output_folder.insert(0, path)
    
    def select_model(self, entry_widget):
        path = filedialog.askopenfilename(
            title="Selecione o modelo DOCX",
            filetypes=[("Word Files", "*.docx"), ("All files", "*.*")]
        )
        if path:
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, path)
    
    def schedule_tasks(self):
        """Agenda todas as tarefas configuradas"""
        excel_path = self.excel_path.get()
        output_folder = self.output_folder.get()
        
        if not excel_path or not output_folder:
            messagebox.showerror("Erro", "Planilha e pasta de saída são obrigatórias")
            return
        
        for entry in self.model_entries:
            modelo_path = entry['modelo_path'].get()
            if not modelo_path:
                continue  # Pula modelos vazios
            
            self.task_manager.agendar_declaracao(
                excel_path=excel_path,
                modelo_path=modelo_path,
                pasta_saida=output_folder,
                filename_pattern=entry['filename_pattern'].get(),
                save_in_subfolders=entry['subfolder_var'].get()
            )
        
        messagebox.showinfo("Sucesso", f"{len(self.model_entries)} modelos agendados!")
        self.destroy()
    
class TemplatePreviewDialog(tk.Toplevel):
    def __init__(self, parent, template_path, sample_data):
        super().__init__(parent)
        self.title("Pré-visualização com Dados Reais")
        
        # Renderiza o DOCX como imagem (simplificado)
        self.canvas = tk.Canvas(self, width=600, height=800)
        self.canvas.pack()
        
        self.render_preview(template_path, sample_data)
    
    def render_preview(self, path, data):
        # 1. Cria uma versão temporária com dados reais
        temp_doc = Document(path)
        self.substitute_placeholders(temp_doc, data)
        
        # 2. Converte para PDF e depois para imagem (simplificado)
        # (Na prática, use libraries como docx2pdf + pdf2image)
        img = self.fake_render_for_demo()  # Implemente isso corretamente
        
        # 3. Exibe no canvas
        self.tk_img = ImageTk.PhotoImage(img)
        self.canvas.create_image(0, 0, anchor="nw", image=self.tk_img)
    
    def substitute_placeholders(self, doc, data):
        for p in doc.paragraphs:
            for key, val in data.items():
                p.text = p.text.replace(f"{{{{{key}}}}}", str(val))

class TemplateManager:
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = Path(templates_dir)
        self.config_path = self.templates_dir / "_config.json"
        self.templates = self._carregar_config()
        # Create templates directory if it doesn't exist
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def _carregar_config(self) -> Dict:
        """Carrega a configuração dos templates"""
        if not self.config_path.exists():
            return {"templates": {}}
        
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, FileNotFoundError):
            return {"templates": {}}

    def _salvar_config(self):
        """Salva as alterações no arquivo de configuração"""
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(self.templates, f, indent=4, ensure_ascii=False)

    def adicionar_template(
        self,
        nome: str,
        arquivo: str,
        categoria: str,
        tags: List[str],
        campos_obrigatorios: List[str],
        descricao: str = ""
    ):
        """Adiciona um novo template à biblioteca"""
        # Convert to Path object if not already
        arquivo_path = Path(arquivo)
        
        # If the file isn't already in the templates directory, copy it there
        if arquivo_path.parent != self.templates_dir:
            new_path = self.templates_dir / arquivo_path.name
            shutil.copy2(arquivo_path, new_path)
            rel_path = arquivo_path.name
        else:
            rel_path = str(arquivo_path.relative_to(self.templates_dir))
        
        self.templates["templates"][nome] = {
            "categoria": categoria,
            "arquivo": rel_path,
            "tags": tags,
            "campos_obrigatorios": campos_obrigatorios,
            "descricao": descricao
        }
        self._salvar_config()

    def buscar_templates(self, categoria=None, tag=None, termo=None):
        """Busca templates com filtros"""
        resultados = []
        
        for nome, template in self.templates["templates"].items():
            # Aplica filtros
            if categoria and template.get("categoria") != categoria:
                continue
                
            if tag and tag not in template.get("tags", []):
                continue
                
            if termo and termo.lower() not in nome.lower():
                continue
                
            # Adiciona os dados do template ao resultado
            resultados.append({
                "nome": nome,
                "categoria": template.get("categoria", ""),
                "tags": template.get("tags", []),
                "campos_obrigatorios": template.get("campos_obrigatorios", []),
                "descricao": template.get("descricao", "")
            })
            
        return resultados

    def get_caminho_template(self, nome_template: str) -> Path:
        """Retorna o caminho absoluto do arquivo do template"""
        if nome_template not in self.templates["templates"]:
            raise ValueError(f"Template '{nome_template}' não encontrado")
            
        return self.templates_dir / self.templates["templates"][nome_template]["arquivo"]

    def validar_dados(self, template_name: str, dados: dict) -> tuple:
        """
        Valida se os dados contêm todos os campos obrigatórios
        Retorna: (bool, lista_de_campos_faltantes)
        """
        template = self.templates["templates"].get(template_name)
        if not template:
            return False, ["Template não encontrado"]
        
        faltantes = [
            campo for campo in template["campos_obrigatorios"]
            if campo not in dados
        ]
        
        return (not faltantes), faltantes

    def get_all_templates(self) -> Dict[str, Dict]:
        """Retorna todos os templates"""
        return self.templates["templates"]

class TemplateBackup:
    def __init__(self, template_manager):
        self.tm = template_manager
        self.backup_dir = Path("backups/templates")
        # Cria todos os diretórios necessários na hierarquia
        self.backup_dir.parent.mkdir(parents=True, exist_ok=True)
        self.backup_dir.mkdir(exist_ok=True)

    def create_backup(self):
        """Cria um backup com timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_file = self.backup_dir / f"templates_{timestamp}.json"
        
        try:
            with open(backup_file, 'w', encoding='utf-8') as f:
                json.dump(self.tm.templates, f, indent=4, ensure_ascii=False)
            return backup_file
        except Exception as e:
            logger.error(f"Error creating backup: {str(e)}")
            raise

class BatchProcessor:
    def __init__(self, max_workers=4):
        self.task_queue = queue.Queue()
        self.executor = ThreadPoolExecutor(max_workers=max_workers)

    def add_task(self, template_path, data):
        self.task_queue.put((template_path, data))

    def process_batch(self):
        futures = []
        while not self.task_queue.empty():
            template, data = self.task_queue.get()
            futures.append(
                self.executor.submit(
                    self._safe_process, 
                    template, 
                    data
                )
            )
        return futures

    def _safe_process(self, template_path, data):
        try:
            from win32com.client import Dispatch
            office_app = Dispatch('Excel.Application')
            doc = office_app.Workbooks.Open(template_path)
            # ... processamento ...
            doc.Save()
            return True
        except Exception as e:
            logger.error(f"Erro em {template_path}: {str(e)}")
            return False
        finally:
            office_app.Quit()

class ResourceMonitor:
    def __init__(self):
        self.max_memory = 0

    def check_resources(self):
        mem = psutil.virtual_memory()
        self.max_memory = max(self.max_memory, mem.used)
        
        if mem.percent > 90:
            raise MemoryError("Uso de memória crítico")
        
        return {
            "memory": f"{mem.used/1024/1024:.2f}MB",
            "cpu": f"{psutil.cpu_percent()}%"
        }


if __name__ == "__main__":
    pythoncom.CoInitialize()
    
    try:
        root = tk.Tk()
        app = DocumentConverterApp(root)
        root.protocol("WM_DELETE_WINDOW", app.on_closing)
        
        if getattr(sys, 'frozen', False):
            app_dir = os.path.dirname(sys.executable)
        else:
            app_dir = os.path.dirname(os.path.abspath(__file__))
        
        icon_path = os.path.join(app_dir, "icon.ico")
        if os.path.exists(icon_path):
            root.iconbitmap(icon_path)
        
        root.mainloop()
    except Exception as e:
        logging.exception("Fatal application error")
        messagebox.showerror("Fatal Error", f"An unexpected error occurred:\n{str(e)}")
    finally:
        pythoncom.CoUninitialize()
